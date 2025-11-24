#!/usr/bin/env python3
# =============================================================================
# VAPT Toolkit - Vulnerability Assessment and Penetration Testing Toolkit
# =============================================================================
#
# Author: Keith Pachulski
# Company: Red Cell Security, LLC
# Email: keith@redcellsecurity.org
# Website: www.redcellsecurity.org
#
# Copyright (c) 2025 Keith Pachulski. All rights reserved.
#
# License: This software is licensed under the MIT License.
#          You are free to use, modify, and distribute this software
#          in accordance with the terms of the license.
#
# Purpose: This script provides an automated installation and management system
#          for a vulnerability assessment and penetration testing
#          toolkit. It installs and configures security tools across multiple
#          categories including exploitation, web testing, network scanning,
#          mobile security, cloud security, and Active Directory testing.
#
# DISCLAIMER: This software is provided "as-is," without warranty of any kind,
#             express or implied, including but not limited to the warranties
#             of merchantability, fitness for a particular purpose, and non-infringement.
#             In no event shall the authors or copyright holders be liable for any claim,
#             damages, or other liability, whether in an action of contract, tort, or otherwise,
#             arising from, out of, or in connection with the software or the use or other dealings
#             in the software.
#
# NOTICE: This toolkit is intended for authorized security testing only.
#         Users are responsible for ensuring compliance with all applicable laws
#         and regulations. Unauthorized use of these tools may violate local,
#         state, federal, and international laws.
#
# =============================================================================

import argparse
import re
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict
import sys


class Finding:
    """Represents a security finding"""

    def __init__(self, title, severity, description, affected_systems,
                 remediation, references, notes, tool_source):
        self.title = title
        self.severity = severity
        self.description = description
        self.affected_systems = affected_systems if isinstance(affected_systems, list) else [affected_systems]
        self.remediation = remediation
        self.references = references if isinstance(references, list) else [references]
        self.notes = notes
        self.tool_source = tool_source

    def merge_with(self, other_finding):
        """Merge another finding of the same type into this one"""
        # Add affected systems if not already present
        for system in other_finding.affected_systems:
            if system not in self.affected_systems:
                self.affected_systems.append(system)

        # DO NOT merge notes - keep only the original proof of existence

        # Merge references
        for ref in other_finding.references:
            if ref and ref not in self.references:
                self.references.append(ref)


class ZAPParser:
    """Parser for OWASP ZAP HTML reports"""

    SEVERITY_MAP = {
        'High': 'High',
        'Medium': 'Medium',
        'Low': 'Low',
        'Informational': 'Informational'
    }

    def __init__(self, html_file):
        self.html_file = html_file
        with open(html_file, 'r', encoding='utf-8') as f:
            self.soup = BeautifulSoup(f.read(), 'html.parser')

    def parse(self):
        """Parse ZAP report and return list of Finding objects"""
        findings = []

        # Find all alert sections
        alert_sections = self.soup.find_all('li', class_='alerts--site-li')

        for site_section in alert_sections:
            site_name = site_section.find('h4')
            if site_name:
                site_name = site_name.get_text(strip=True).split('(')[0].strip()

            # Find all alert types within this site - using direct children only
            alert_type_lis = []
            for child_ol in site_section.find_all('ol', recursive=False):
                for li in child_ol.find_all('li', recursive=False):
                    alert_type_lis.append(li)

            for alert_li in alert_type_lis:
                # Get alert name from h5
                h5 = alert_li.find('h5', recursive=False)
                if not h5:
                    continue

                alert_link = h5.find('a')
                if not alert_link:
                    continue

                alert_name = alert_link.get_text(strip=True)

                # Get all instances of this alert - look in nested ol > li > details
                instance_lis = []
                for ol in alert_li.find_all('ol', recursive=False):
                    instance_lis.extend(ol.find_all('li', recursive=False))

                for instance_li in instance_lis:
                    # Get the details tag
                    detail_item = instance_li.find('details', recursive=False)
                    if not detail_item:
                        continue

                    # Get URL from summary
                    summary = detail_item.find('summary')
                    url = ""
                    if summary:
                        url_span = summary.find('span', class_='request-method-n-url')
                        if url_span:
                            url = url_span.get_text(strip=True)

                    # Find the table with alert details
                    alert_table = detail_item.find('table', class_='alerts-table')
                    if not alert_table:
                        continue

                    # Extract data from table rows
                    description = ""
                    solution = ""
                    references = []
                    request = ""
                    response = ""
                    evidence = ""

                    rows = alert_table.find_all('tr')
                    for row in rows:
                        header = row.find('th')
                        if not header:
                            continue

                        header_text = header.get_text(strip=True)
                        td = row.find('td')
                        if not td:
                            continue

                        if header_text == 'Alert description':
                            description = td.get_text(strip=True)
                        elif header_text == 'Solution':
                            solution = td.get_text(strip=True)
                        elif header_text == 'Alert tags':
                            # Extract references from links
                            links = td.find_all('a')
                            for link in links:
                                href = link.get('href', '')
                                if href and ('cwe.mitre.org' in href or 'owasp.org' in href or 'zaproxy.org' in href):
                                    references.append(href)
                        elif header_text == 'Request':
                            request_details = td.find('details')
                            if request_details:
                                code = request_details.find('code')
                                if code:
                                    request = code.get_text()[:500]  # Limit length
                        elif header_text == 'Response':
                            response_details = td.find('details')
                            if response_details:
                                code = response_details.find('code')
                                if code:
                                    response = code.get_text()[:500]  # Limit length
                        elif header_text == 'Evidence':
                            evidence = td.get_text(strip=True)

                    # Determine severity from parent section
                    severity = self._extract_severity(site_section)

                    # Build notes section
                    notes = f"Tool: OWASP ZAP\nURL: {url}"
                    if evidence:
                        notes += f"\n\nEvidence: {evidence}"
                    if request:
                        notes += f"\n\nRequest Sample:\n{request}"
                    if response:
                        notes += f"\n\nResponse Sample:\n{response}"

                    finding = Finding(
                        title=alert_name,
                        severity=severity,
                        description=description,
                        affected_systems=url,
                        remediation=solution,
                        references=references,
                        notes=notes,
                        tool_source='ZAP'
                    )

                    findings.append(finding)

        return findings

    def _extract_severity(self, element):
        """Extract severity from parent sections"""
        # Look for risk level in parent sections
        parent = element
        while parent:
            if parent.name == 'li' and parent.get('id', '').startswith('alerts--risk-'):
                risk_id = parent.get('id', '')
                # Extract risk level from id (e.g., alerts--risk-2-confidence-3)
                if 'risk-3' in risk_id:
                    return 'High'
                elif 'risk-2' in risk_id:
                    return 'Medium'
                elif 'risk-1' in risk_id:
                    return 'Low'
                elif 'risk-0' in risk_id:
                    return 'Informational'
            parent = parent.parent
        return 'Medium'  # Default


class BurpParser:
    """Parser for Burp Suite HTML reports"""

    SEVERITY_MAP = {
        'High': 'High',
        'Medium': 'Medium',
        'Low': 'Low',
        'Information': 'Informational'
    }

    def __init__(self, html_file):
        self.html_file = html_file
        with open(html_file, 'r', encoding='utf-8') as f:
            self.soup = BeautifulSoup(f.read(), 'html.parser')

    def parse(self):
        """Parse Burp report and return list of Finding objects"""
        findings = []

        # Find all issue sections (marked by BODH0 class for main issues)
        issue_sections = self.soup.find_all('span', class_='BODH0')

        for issue_section in issue_sections:
            issue_link = issue_section.find('a')
            if not issue_link:
                continue

            issue_title = issue_link.get_text(strip=True)
            issue_id = issue_section.get('id', '')

            # Gather general background and remediation
            general_description = ""
            general_remediation = ""
            general_references = []

            # Find sections belonging to this issue
            next_sibling = issue_section.find_next_sibling()
            instances = []

            while next_sibling:
                if isinstance(next_sibling, str):
                    next_sibling = next_sibling.find_next_sibling()
                    continue

                if next_sibling.name == 'span' and 'BODH0' in next_sibling.get('class', []):
                    break  # reached next issue

                if next_sibling.name == 'h2':
                    header_text = next_sibling.get_text(strip=True)
                    content_elem = next_sibling.find_next_sibling()
                    if content_elem and content_elem.name == 'span':
                        content = content_elem.get_text(strip=True)
                        if header_text == 'Issue background':
                            general_description = content
                        elif header_text == 'Issue remediation':
                            general_remediation = content
                        elif header_text == 'References':
                            links = content_elem.find_all('a')
                            for link in links:
                                href = link.get('href', '')
                                if href:
                                    general_references.append(href)

                # Capture instances (sub-findings)
                if next_sibling.name == 'span' and 'BODH1' in next_sibling.get('class', []):
                    instances.append(next_sibling)

                next_sibling = next_sibling.find_next_sibling()

            # If no sub-instances, create one from summary
            if not instances:
                severity = "Medium"
                host = ""
                path = ""

                summary_table = issue_section.find_next('table', class_='summary_table')
                if summary_table:
                    rows = summary_table.find_all('tr')
                    for row in rows:
                        tds = row.find_all('td')
                        if len(tds) >= 2:
                            label = tds[0].get_text(strip=True).rstrip(':')
                            value = tds[1].get_text(strip=True)
                            if label == 'Severity':
                                severity = self.SEVERITY_MAP.get(value, value)
                            elif label == 'Host':
                                host = value
                            elif label == 'Path':
                                path = value

                affected_system = f"{host}{path}" if host else "See report for details"
                finding = Finding(
                    title=issue_title,
                    severity=severity,
                    description=general_description,
                    affected_systems=affected_system,
                    remediation=general_remediation,
                    references=general_references,
                    notes=f"Tool: Burp Suite",
                    tool_source='Burp'
                )
                findings.append(finding)
                continue

            # Process each sub-instance
            for instance in instances:
                instance_text = instance.get_text(strip=True)
                instance_url = instance_text.split('\xa0', 1)[-1] if '\xa0' in instance_text else instance_text

                summary_table = instance.find_next('table', class_='summary_table')

                severity = "Medium"
                host = ""
                path = ""

                if summary_table:
                    rows = summary_table.find_all('tr')
                    for row in rows:
                        tds = row.find_all('td')
                        if len(tds) >= 2:
                            label = tds[0].get_text(strip=True).rstrip(':')
                            value = tds[1].get_text(strip=True)
                            if label == 'Severity':
                                severity = self.SEVERITY_MAP.get(value, value)
                            elif label == 'Host':
                                host = value
                            elif label == 'Path':
                                path = value

                instance_description = general_description
                instance_detail = ""

                # Locate "Issue detail"
                next_elem = summary_table.find_next_sibling() if summary_table else instance.find_next_sibling()
                while next_elem:
                    if isinstance(next_elem, str):
                        next_elem = next_elem.find_next_sibling()
                        continue
                    if next_elem.name == 'span' and ('BODH0' in next_elem.get('class', []) or 'BODH1' in next_elem.get('class', [])):
                        break
                    if next_elem.name == 'h2' and 'Issue detail' in next_elem.get_text():
                        detail_elem = next_elem.find_next_sibling()
                        if detail_elem and detail_elem.name == 'span':
                            instance_detail = detail_elem.get_text(strip=True)[:500]
                    next_elem = next_elem.find_next_sibling()
                    if instance_detail:
                        break

                # ✅ NEW SECTION: Extract Request/Response
                request = response = ""
                rr_divs = []
                next_rr = next_elem
                while next_rr:
                    if isinstance(next_rr, str):
                        next_rr = next_rr.find_next_sibling()
                        continue
                    if next_rr.name == 'div' and 'rr_div' in next_rr.get('class', []):
                        rr_divs.append(next_rr)
                    elif next_rr.name == 'span' and ('BODH0' in next_rr.get('class', []) or 'BODH1' in next_rr.get('class', [])):
                        break
                    next_rr = next_rr.find_next_sibling()

                if len(rr_divs) >= 1:
                    request = rr_divs[0].get_text(strip=True)[:500]
                if len(rr_divs) >= 2:
                    response = rr_divs[1].get_text(strip=True)[:500]

                affected_system = f"{host}{path}" if host else instance_url

                # Build notes including request/response
                notes = f"Tool: Burp Suite\nURL: {affected_system}"
                if instance_detail:
                    notes += f"\n\nDetails:\n{instance_detail}"
                if request:
                    notes += f"\n\nRequest Sample:\n{request}"
                if response:
                    notes += f"\n\nResponse Sample:\n{response}"

                finding = Finding(
                    title=issue_title,
                    severity=severity,
                    description=instance_description if instance_description else instance_detail,
                    affected_systems=affected_system,
                    remediation=general_remediation,
                    references=general_references,
                    notes=notes,
                    tool_source='Burp'
                )

                findings.append(finding)

        return findings

class ReportGenerator:
    """Generates Word document with findings"""

    SEVERITY_ORDER = {
        'Critical': 0,
        'High': 1,
        'Medium': 2,
        'Low': 3,
        'Informational': 4
    }

    def __init__(self, findings):
        self.findings = findings
        self.doc = Document()

    def deduplicate_findings(self):
        """Deduplicate findings by title and merge affected systems"""
        finding_dict = {}

        for finding in self.findings:
            key = finding.title.lower().strip()

            if key in finding_dict:
                # Merge with existing finding
                finding_dict[key].merge_with(finding)
            else:
                finding_dict[key] = finding

        return list(finding_dict.values())

    def sort_findings(self, findings):
        """Sort findings by severity"""
        return sorted(findings,
                     key=lambda f: self.SEVERITY_ORDER.get(f.severity, 99))

    def generate(self, output_file):
        """Generate Word document with findings matching Nessus report format"""
        # Filter out informational findings
        filtered_findings = [f for f in self.findings if f.severity.lower() != 'informational']
        informational_count = len(self.findings) - len(filtered_findings)

        # Deduplicate findings
        unique_findings = self._deduplicate_findings(filtered_findings)

        # Sort by severity
        severity_order = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3}
        sorted_findings = sorted(unique_findings, key=lambda x: (
            severity_order.get(x.severity, 999),
            x.title.lower()
        ))

        # Organize by severity
        findings_by_severity = defaultdict(list)
        for finding in sorted_findings:
            findings_by_severity[finding.severity].append(finding)

        # Severity colors
        severity_colors = {
            'Critical': RGBColor(139, 0, 0),    # Dark Red
            'High': RGBColor(255, 0, 0),        # Red
            'Medium': RGBColor(255, 140, 0),    # Dark Orange
            'Low': RGBColor(65, 105, 225)       # Royal Blue
        }

        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # Title
        title = self.doc.add_heading('Web Application Security Assessment Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Timestamp
        from datetime import datetime
        timestamp = self.doc.add_paragraph()
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = timestamp.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(102, 102, 102)

        # Executive Summary
        self.doc.add_heading('Executive Summary', level=1)

        # Summary statistics
        summary = self.doc.add_paragraph(
            f'This report contains {len(sorted_findings)} unique security findings organized by severity:'
        )

        # Severity breakdown
        for severity in ['Critical', 'High', 'Medium', 'Low']:
            count = len(findings_by_severity[severity])
            if count > 0:
                p = self.doc.add_paragraph(style='List Bullet')
                run = p.add_run(f'{severity}: {count}')
                run.bold = True
                run.font.color.rgb = severity_colors.get(severity, RGBColor(0, 0, 0))

        # Findings section
        self.doc.add_heading('Findings', level=1)

        # Process each severity level
        for severity in ['Critical', 'High', 'Medium', 'Low']:
            findings = findings_by_severity[severity]

            if not findings:
                continue

            # Severity heading (H2)
            severity_heading = self.doc.add_heading(severity, level=2)
            severity_heading.runs[0].font.color.rgb = severity_colors.get(severity, RGBColor(0, 0, 0))

            # Individual findings
            for finding in findings:
                # Finding heading (H3) - Name | Severity
                heading = self.doc.add_heading(level=3)
                heading.add_run(finding.title).bold = True
                heading.add_run(' | ').bold = True
                run = heading.add_run(finding.severity)
                run.bold = True
                run.font.color.rgb = severity_colors.get(finding.severity, RGBColor(0, 0, 0))

                # Affected Systems
                label = self.doc.add_paragraph()
                label.add_run('Affected System(s):').bold = True

                if finding.affected_systems and finding.affected_systems != ['N/A']:
                    for system in finding.affected_systems:
                        self.doc.add_paragraph(system, style='List Bullet')
                else:
                    p = self.doc.add_paragraph('None')
                    p.runs[0].italic = True
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

                # Description
                label = self.doc.add_paragraph()
                label.add_run('Description:').bold = True
                if finding.description:
                    self.doc.add_paragraph(finding.description)
                else:
                    p = self.doc.add_paragraph('N/A')
                    p.runs[0].italic = True
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

                # Remediation
                label = self.doc.add_paragraph()
                label.add_run('Remediation:').bold = True
                if finding.remediation:
                    self.doc.add_paragraph(finding.remediation)
                else:
                    p = self.doc.add_paragraph('See references for remediation guidance.')
                    p.runs[0].italic = True
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

                # References
                label = self.doc.add_paragraph()
                label.add_run('References:').bold = True
                if finding.references and finding.references != ['N/A']:
                    for ref in finding.references:
                        if ref and ref != 'N/A':
                            self.doc.add_paragraph(ref)
                else:
                    p = self.doc.add_paragraph('None')
                    p.runs[0].italic = True
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

                # Evidence/Notes
                label = self.doc.add_paragraph()
                label.add_run('Evidence:').bold = True
                if finding.notes:
                    self.doc.add_paragraph(finding.notes)
                else:
                    p = self.doc.add_paragraph('N/A')
                    p.runs[0].italic = True
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

                # Spacing between findings
                self.doc.add_paragraph()

        # Save document
        self.doc.save(output_file)
        print(f"\n[+] Report generated successfully: {output_file}")
        print(f"[+] Total unique findings: {len(sorted_findings)}")
        print(f"[+] Original findings before deduplication: {len(filtered_findings)}")
        if informational_count > 0:
            print(f"[+] Informational findings excluded: {informational_count}")

def find_reports_in_directory(directory):
    """
    Scan directory for ZAP and Burp HTML reports
    Returns: (list of zap files, list of burp files)
    """
    directory = Path(directory)

    if not directory.exists():
        print(f"[!] Error: Directory not found: {directory}")
        sys.exit(1)

    if not directory.is_dir():
        print(f"[!] Error: Path is not a directory: {directory}")
        sys.exit(1)

    zap_reports = []
    burp_reports = []

    # Get all HTML files in directory
    html_files = list(directory.glob("*.html")) + list(directory.glob("*.htm"))

    print(f"[*] Scanning directory: {directory}")
    print(f"[*] Found {len(html_files)} HTML file(s)")

    for html_file in html_files:
        try:
            with open(html_file, 'r', encoding='utf-8') as f:
                content = f.read(5000)  # Read first 5000 chars to identify

                # Check for ZAP signatures
                if 'ZAP by Checkmarx' in content or 'ZAP Scanning Report' in content or 'zaproxy.org' in content:
                    zap_reports.append(str(html_file))
                    print(f"  [ZAP] {html_file.name}")

                # Check for Burp signatures
                elif 'Burp Scanner Report' in content or 'Burp Suite' in content or 'portswigger' in content:
                    burp_reports.append(str(html_file))
                    print(f"  [BURP] {html_file.name}")

                else:
                    print(f"  [SKIP] {html_file.name} - Unknown format")

        except Exception as e:
            print(f"  [ERROR] {html_file.name} - Could not read: {e}")

    return zap_reports, burp_reports

def display_test_finding(findings):
    """Display the first parsed finding for testing."""
    if not findings:
        print("\n[!] No findings found to display")
        return

    print("\n" + "="*80)
    print("TEST MODE - Displaying First Parsed Finding")
    print("="*80 + "\n")

    finding = findings[0]

    print(f"Severity: {finding.severity}")
    print(f"Finding: {finding.title}")
    print(f"Tool Source: {finding.tool_source}")

    print(f"\nAffected System(s):")
    for system in finding.affected_systems:
        print(f"  - {system}")

    print(f"\nDescription:")
    desc_preview = finding.description[:500] + "..." if len(finding.description) > 500 else finding.description
    print(desc_preview)

    print(f"\nRemediation:")
    rem_preview = finding.remediation[:500] + "..." if len(finding.remediation) > 500 else finding.remediation
    print(rem_preview)

    if finding.references:
        print(f"\nReferences:")
        for ref in finding.references[:10]:  # Limit to first 10
            print(f"  {ref}")
        if len(finding.references) > 10:
            print(f"  ... and {len(finding.references) - 10} more")

    if finding.notes:
        print(f"\nEvidence/Notes:")
        notes_preview = finding.notes[:500] + "..." if len(finding.notes) > 500 else finding.notes
        print(notes_preview)

    print("\n" + "="*80)

def main():
    parser = argparse.ArgumentParser(
        description='Parse OWASP ZAP and Burp Suite HTML reports and generate Word document',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process all reports in a directory (auto-detects ZAP and Burp)
  python burp_zap_parser.py -d ./scan_reports -o findings

  # Test mode - preview first finding without generating report
  python burp_zap_parser.py -d ./scan_reports --test

  # Process specific files manually
  python burp_zap_parser.py -z report1.html report2.html -o findings
  python burp_zap_parser.py -b burp1.html burp2.html -o findings
  python burp_zap_parser.py -z zap1.html -b burp1.html -o findings

Note: Informational findings are automatically excluded from the output.
      Duplicate findings across multiple files are merged with all affected systems listed.
        """
    )

    parser.add_argument('-d', '--directory', type=str,
                       help='Directory containing ZAP and/or Burp HTML reports (auto-detects report types)')
    parser.add_argument('-z', '--zap', type=str, nargs='+',
                       help='Path(s) to ZAP HTML report(s) - can specify multiple files')
    parser.add_argument('-b', '--burp', type=str, nargs='+',
                       help='Path(s) to Burp Suite HTML report(s) - can specify multiple files')
    parser.add_argument('-o', '--output', type=str,
                       help='Output Word document path')
    parser.add_argument('-t', '--test', action='store_true',
                       help='Test mode: display first parsed finding without generating report')

    args = parser.parse_args()

    # Validate inputs
    if not args.directory and not args.zap and not args.burp:
        print("[!] Error: Must specify either -d (directory) OR -z/-b (specific files)")
        parser.print_help()
        sys.exit(1)

    # Cannot use both directory and specific files
    if args.directory and (args.zap or args.burp):
        print("[!] Error: Cannot use -d (directory) with -z/-b (specific files). Choose one approach.")
        sys.exit(1)

    # Output required unless in test mode
    if not args.test and not args.output:
        print("[!] Error: -o (output) is required unless using --test mode")
        parser.print_help()
        sys.exit(1)

    # Ensure output file has .docx extension (if not in test mode)
    if args.output and not args.output.lower().endswith('.docx'):
        args.output = f"{args.output}.docx"
        print(f"[*] Output filename adjusted to: {args.output}")

    # If directory mode, auto-detect reports
    if args.directory:
        zap_files, burp_files = find_reports_in_directory(args.directory)

        if not zap_files and not burp_files:
            print("[!] Error: No ZAP or Burp reports found in directory")
            sys.exit(1)

        print(f"\n[+] Identified {len(zap_files)} ZAP report(s) and {len(burp_files)} Burp report(s)")
    else:
        # Manual file mode
        zap_files = args.zap if args.zap else []
        burp_files = args.burp if args.burp else []

    all_findings = []

    # Parse ZAP report(s)
    if zap_files:
        print(f"\n[*] Processing {len(zap_files)} ZAP report(s)...")
        for zap_file in zap_files:
            if not Path(zap_file).exists():
                print(f"[!] Error: ZAP report not found: {zap_file}")
                sys.exit(1)

            print(f"  [*] Parsing: {Path(zap_file).name}")
            zap_parser = ZAPParser(zap_file)
            zap_findings = zap_parser.parse()
            all_findings.extend(zap_findings)
            print(f"      Found {len(zap_findings)} findings")

    # Parse Burp report(s)
    if burp_files:
        print(f"\n[*] Processing {len(burp_files)} Burp report(s)...")
        for burp_file in burp_files:
            if not Path(burp_file).exists():
                print(f"[!] Error: Burp report not found: {burp_file}")
                sys.exit(1)

            print(f"  [*] Parsing: {Path(burp_file).name}")
            burp_parser = BurpParser(burp_file)
            burp_findings = burp_parser.parse()
            all_findings.extend(burp_findings)
            print(f"      Found {len(burp_findings)} findings")

    if not all_findings:
        print("\n[!] Warning: No findings were extracted from the reports")
        sys.exit(0)

    print(f"\n[*] Total findings collected: {len(all_findings)}")

    # Filter informational and sort
    filtered_findings = [f for f in all_findings if f.severity.lower() != 'informational']
    informational_count = len(all_findings) - len(filtered_findings)

    if informational_count > 0:
        print(f"[*] Informational findings excluded: {informational_count}")

    # Sort by severity
    severity_order = {'Critical': 0, 'High': 1, 'Medium': 2, 'Low': 3}
    sorted_findings = sorted(filtered_findings, key=lambda x: (
        severity_order.get(x.severity, 999),
        x.title.lower()
    ))

    # Test mode
    if args.test:
        display_test_finding(sorted_findings)
        sys.exit(0)

    # Generate report
    print(f"\n[*] Generating Word document...")
    generator = ReportGenerator(all_findings)
    generator.generate(args.output)

    print(f"\n[+] Done! You can now copy/paste findings from {args.output} into your formal report.")

if __name__ == '__main__':
    main()
