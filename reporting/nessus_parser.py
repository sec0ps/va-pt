#!/usr/bin/env python3
# =============================================================================
# VAPT Toolkit - Vulnerability Assessment and Penetration Testing Toolkit
# =============================================================================
#
# Author: Keith Pachulski
# Company: Red Cell Security, LLC
# Email: keith@redcellsecurity.org
# Website: www.redcellsecurity.org
#
# Copyright (c) 2026 Keith Pachulski. All rights reserved.
#
# License: This software is licensed under the MIT License.
#          You are free to use, modify, and distribute this software
#          in accordance with the terms of the license.
#
# Purpose: Parses Nessus (.nessus / NessusClientData_v2 XML) scan output and
#          generates a consolidated DOCX assessment report. Scores severity
#          from the authoritative severity attribute, deduplicates findings by
#          plugin across hosts while preserving per-target evidence, condenses
#          version-ladder plugin cascades (e.g. dozens of "Apache Tomcat X <
#          Y" gates) into a single validated finding, and optionally merges
#          multiple .nessus files into one. Report layout matches the unified
#          VAPT report format emitted by vapt_report_parser.py.
#
# DISCLAIMER: This software is provided "as-is," without warranty of any kind,
#             express or implied, including but not limited to the warranties
#             of merchantability, fitness for a particular purpose, and non-infringement.
#             In no event shall the authors or copyright holders be liable for any claim,
#             damages, or other liability, whether in an action of contract, tort, or otherwise,
#             arising from, out of, or in connection with the software or the use or other dealings
#             in the software.
#
# NOTICE: This toolkit is intended for authorized security testing only.
#         Users are responsible for ensuring compliance with all applicable laws
#         and regulations. Unauthorized use of these tools may violate local,
#         state, federal, and international laws.
#
# =============================================================================

"""
Nessus parser and DOCX report generator for the VAPT toolkit.

Severity is taken from the ReportItem ``severity`` attribute (0=Info, 1=Low,
2=Medium, 3=High, 4=Critical) rather than the legacy ``risk_factor`` field,
which is frequently stale relative to the CVSS-derived attribute. Findings are
deduplicated per plugin across all hosts and ports, retaining each affected
target's ``plugin_output`` as per-target evidence.

Nessus fires a distinct plugin for every fixed-version gate above an installed
version, so an outdated component surfaces as dozens of near-identical
findings across several severity buckets. Findings whose name carries a
``< <version>`` ladder marker are grouped by normalized product name and
condensed into one finding: maximum severity, the union of CVEs and external
references, remediation pointing at the highest fixed version observed, and
per-host installed-vs-fixed evidence so the consolidation stays verifiable.
Non-ladder findings (SSL certificate issues, TLS version detections, default
files, and so on) are never merged.
"""

import os
import re
import sys
import signal
import argparse
import xml.etree.ElementTree as ET
from pathlib import Path
from datetime import datetime
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SEVERITY_LABELS = {'3': 'High', '2': 'Medium', '1': 'Low'}
SEVERITY_ORDER = ['High', 'Medium', 'Low']

LADDER_RE = re.compile(r'<\s*\d')
PRODUCT_SPLIT_RE = re.compile(r'\s+<?\s*\d')
FIXED_IN_NAME_RE = re.compile(r'<\s*([0-9][\w.]+)')
FIXED_IN_OUTPUT_RE = re.compile(r'Fixed\s+version\s*:\s*([0-9][\w.]+)', re.IGNORECASE)
INSTALLED_IN_OUTPUT_RE = re.compile(r'Installed\s+version\s*:\s*([0-9][\w.]+)', re.IGNORECASE)
URL_IN_OUTPUT_RE = re.compile(r'URL\s*:\s*(\S+)', re.IGNORECASE)
HTTP_404_RE = re.compile(r'HTTP/\d(?:\.\d)?\s+404\b')
TOOL_NAME_RE = re.compile(r'\b(?:nessus|tenable)\b', re.IGNORECASE)
TOOL_URL_RE = re.compile(r'nessus\.org|tenable\.com', re.IGNORECASE)


def scrub_tool_names(text):
    if not text:
        return text
    return TOOL_NAME_RE.sub('automated testing', text)


class NessusFinding:

    def __init__(self, plugin_id):
        self.plugin_id = plugin_id
        self.name = ""
        self.severity_int = 0
        self.cvss_score = ""
        self.description = ""
        self.solution = ""
        self.references = []
        self.evidence_by_system = {}
        self.is_ladder = False
        self.is_merged = False
        self.product = ""
        self.fixed_version = ""
        self.evidence_override = None
        self.original_severity_int = 0
        self.exploit_available = False
        self.exploit_metasploit = False
        self.exploit_kev = False
        self.exploit_ease = False

    @property
    def severity(self):
        return SEVERITY_LABELS.get(str(self.severity_int), 'Informational')

    def systems(self):
        return sorted(self.evidence_by_system.keys())

    def has_cve(self):
        return any(r.startswith('CVE:') for r in self.references)

    def exploitable(self):
        return (self.exploit_kev or self.exploit_metasploit
                or self.exploit_available or self.exploit_ease)

    def to_dict(self):
        return {
            'plugin_id': self.plugin_id,
            'name': self.name,
            'severity': self.severity,
            'cvss_score': self.cvss_score,
            'description': self.description,
            'solution': self.solution,
            'references': self.references,
            'affected_systems': self.systems(),
            'evidence_by_system': self.evidence_by_system,
        }


def find_nessus_files(directory):
    nessus_files = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('.nessus'):
                nessus_files.append(os.path.join(root, file))
    return sorted(nessus_files)


def parse_nessus_file(filepath):
    try:
        return ET.parse(filepath)
    except ET.ParseError as e:
        print(f"[!] Parse error in {filepath}: {e}")
        return None
    except Exception as e:
        print(f"[!] Error reading {filepath}: {e}")
        return None


def extract_text(element, tag):
    child = element.find(tag)
    return child.text if child is not None and child.text else ""


def extract_references(report_item):
    refs = []
    for cve in report_item.findall('cve'):
        if cve.text:
            refs.append(f"CVE: {cve.text}")
    for xref in report_item.findall('xref'):
        if xref.text and not TOOL_URL_RE.search(xref.text):
            refs.append(xref.text)
    for see_also in report_item.findall('see_also'):
        if see_also.text:
            for line in see_also.text.splitlines():
                line = line.strip()
                if line and not TOOL_URL_RE.search(line):
                    refs.append(line)
    return refs


def system_string(host_name, report_item):
    port = report_item.get('port', '')
    protocol = report_item.get('protocol', '')
    svc_name = report_item.get('svc_name', '')
    info = host_name
    if port and port != '0':
        info += f":{port}"
    if protocol:
        info += f" ({protocol}"
        if svc_name:
            info += f"/{svc_name}"
        info += ")"
    return info


def clean_output(text):
    if not text:
        return ""
    lines = [ln.rstrip() for ln in text.splitlines()]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    return "\n".join(lines)


def version_key(version):
    parts = re.findall(r'\d+', version or '')
    return tuple(int(p) for p in parts) if parts else (0,)


def product_name(name):
    base = PRODUCT_SPLIT_RE.split(name, 1)[0]
    return base.strip().rstrip('<').strip()


def highest_fixed_version(names, outputs):
    candidates = []
    for name in names:
        candidates.extend(FIXED_IN_NAME_RE.findall(name))
    for output in outputs:
        candidates.extend(FIXED_IN_OUTPUT_RE.findall(output))
    best = ""
    best_key = None
    for version in candidates:
        key = version_key(version)
        if best_key is None or key > best_key:
            best_key = key
            best = version
    return best


def normalized_ladder_output(raw, fixed):
    installed_match = INSTALLED_IN_OUTPUT_RE.search(raw or "")
    if not installed_match:
        return raw
    lines = []
    url_match = URL_IN_OUTPUT_RE.search(raw or "")
    if url_match:
        lines.append(f"URL               : {url_match.group(1)}")
    lines.append(f"Installed version : {installed_match.group(1)}")
    if fixed:
        lines.append(f"Fixed version     : {fixed}")
    return "\n".join(lines)


def dedup_references(references):
    cves = []
    others = []
    seen = set()
    for ref in references:
        if ref in seen:
            continue
        seen.add(ref)
        (cves if ref.startswith('CVE:') else others).append(ref)
    return sorted(cves) + sorted(others)


def condense_ladders(plugin_findings):
    ladder_groups = defaultdict(list)
    final = []

    for finding in plugin_findings:
        if finding.is_ladder:
            ladder_groups[finding.product.lower()].append(finding)
        else:
            final.append(finding)

    for group in ladder_groups.values():
        if len(group) == 1:
            final.append(group[0])
            continue
        final.append(merge_ladder_group(group))

    return final


def merge_ladder_group(group):
    ordered = sorted(
        group,
        key=lambda f: highest_fixed_version([f.name], list(f.evidence_by_system.values())) or "",
        reverse=True,
    )
    product = ordered[0].product
    fixed = highest_fixed_version(
        [f.name for f in group],
        [out for f in group for out in f.evidence_by_system.values()],
    )

    merged = NessusFinding(plugin_id="+".join(f.plugin_id for f in group))
    merged.severity_int = max(f.severity_int for f in group)
    merged.cvss_score = max((f.cvss_score for f in group if f.cvss_score), default="", key=_cvss_key)
    merged.is_ladder = True
    merged.is_merged = True
    merged.product = product
    merged.fixed_version = fixed
    merged.exploit_available = any(f.exploit_available for f in group)
    merged.exploit_metasploit = any(f.exploit_metasploit for f in group)
    merged.exploit_kev = any(f.exploit_kev for f in group)
    merged.exploit_ease = any(f.exploit_ease for f in group)

    if fixed:
        merged.name = f"{product} — Outdated Version (Multiple Vulnerabilities)"
        merged.solution = f"Upgrade {product} to version {fixed} or later."
        merged.description = (
            f"The detected {product} installation is running an outdated version and is "
            f"affected by multiple vulnerabilities addressed across successive releases. "
            f"Automated testing identified {len(group)} separate version-threshold findings "
            f"for this component; they are consolidated here into a single finding. The most "
            f"recent fixed version identified is {fixed}. Per-host installed versions are shown "
            f"under Evidence."
        )
    else:
        merged.name = f"{product} — Outdated Version (Multiple Vulnerabilities)"
        merged.solution = f"Upgrade {product} to the latest supported release."
        merged.description = (
            f"The detected {product} installation is running an outdated version and is "
            f"affected by multiple vulnerabilities addressed across successive releases. "
            f"Automated testing identified {len(group)} separate version-threshold findings "
            f"for this component; they are consolidated here into a single finding. Per-host "
            f"installed versions are shown under Evidence."
        )

    references = []
    for finding in group:
        references.extend(r for r in finding.references if r.startswith('CVE:'))
    merged.references = dedup_references(references)

    for finding in ordered:
        for system, output in finding.evidence_by_system.items():
            if system not in merged.evidence_by_system or not merged.evidence_by_system[system]:
                merged.evidence_by_system[system] = output

    for system in list(merged.evidence_by_system.keys()):
        merged.evidence_by_system[system] = normalized_ladder_output(
            merged.evidence_by_system[system], fixed
        )

    return merged


def _cvss_key(value):
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def exploit_label(finding):
    if not finding.has_cve():
        return None
    if finding.exploitable():
        quals = []
        if finding.exploit_kev:
            quals.append('CISA KEV — actively exploited')
        if finding.exploit_metasploit:
            quals.append('Metasploit')
        if not quals and (finding.exploit_available or finding.exploit_ease):
            quals.append('public exploit')
        suffix = f" ({'; '.join(quals)})" if quals else ""
        return f"Yes{suffix}"
    if finding.original_severity_int != finding.severity_int:
        original = SEVERITY_LABELS.get(str(finding.original_severity_int), '')
        adjusted = SEVERITY_LABELS.get(str(finding.severity_int), '')
        return f"No — severity adjusted from {original} to {adjusted}"
    return "No"


def apply_severity_adjustment(finding):
    finding.original_severity_int = finding.severity_int
    if finding.has_cve() and not finding.exploitable() and finding.severity_int > 1:
        finding.severity_int -= 1
        return True
    return False


def parse_findings(nessus_files):
    findings_map = {}
    raw_item_count = 0
    dropped_404 = 0

    for nessus_file in nessus_files:
        tree = parse_nessus_file(nessus_file)
        if tree is None:
            continue
        report = tree.getroot().find('Report')
        if report is None:
            continue

        for report_host in report.findall('ReportHost'):
            host_name = report_host.get('name', 'Unknown')
            for report_item in report_host.findall('ReportItem'):
                severity_int = int(report_item.get('severity', '0') or '0')
                if severity_int == 0:
                    continue
                severity_int = min(severity_int, 3)

                raw_output = clean_output(extract_text(report_item, 'plugin_output'))
                if HTTP_404_RE.search(raw_output):
                    dropped_404 += 1
                    continue
                output = scrub_tool_names(raw_output)

                raw_item_count += 1

                plugin_id = report_item.get('pluginID', '')
                if plugin_id not in findings_map:
                    finding = NessusFinding(plugin_id)
                    finding.name = report_item.get('pluginName', 'Unknown')
                    finding.severity_int = severity_int
                    finding.cvss_score = (
                        extract_text(report_item, 'cvss3_base_score')
                        or extract_text(report_item, 'cvss_base_score')
                    )
                    finding.description = scrub_tool_names(extract_text(report_item, 'description'))
                    finding.solution = scrub_tool_names(extract_text(report_item, 'solution'))
                    finding.references = dedup_references(extract_references(report_item))
                    finding.is_ladder = bool(LADDER_RE.search(finding.name))
                    finding.product = product_name(finding.name) if finding.is_ladder else ""
                    findings_map[plugin_id] = finding

                finding = findings_map[plugin_id]
                system = system_string(host_name, report_item)
                if system not in finding.evidence_by_system or not finding.evidence_by_system[system]:
                    finding.evidence_by_system[system] = output

                if extract_text(report_item, 'exploit_available').strip().lower() == 'true':
                    finding.exploit_available = True
                if report_item.find('exploit_framework_metasploit') is not None:
                    finding.exploit_metasploit = True
                if any(x.text and 'CISA' in x.text.upper() for x in report_item.findall('xref')):
                    finding.exploit_kev = True
                ease = extract_text(report_item, 'exploitability_ease').strip().lower()
                if ease and 'no known exploits' not in ease:
                    finding.exploit_ease = True

    unique_plugins = len(findings_map)
    condensed = condense_ladders(list(findings_map.values()))

    adjusted_count = 0
    for finding in condensed:
        if apply_severity_adjustment(finding):
            adjusted_count += 1

    organized = {sev: [] for sev in SEVERITY_ORDER}
    for finding in condensed:
        if finding.severity in organized:
            organized[finding.severity].append(finding)

    for sev in SEVERITY_ORDER:
        organized[sev].sort(key=lambda f: f.name.lower())

    return organized, raw_item_count, unique_plugins, dropped_404, adjusted_count


def merge_nessus_files(nessus_files, output_file):
    if not nessus_files:
        print("[!] No .nessus files found to merge")
        return False

    print(f"[*] Found {len(nessus_files)} .nessus file(s)")

    base_tree = parse_nessus_file(nessus_files[0])
    if base_tree is None:
        print(f"[!] Failed to parse base file: {nessus_files[0]}")
        return False

    base_root = base_tree.getroot()
    base_report = base_root.find('Report')
    if base_report is None:
        print("[!] No Report element found in base file")
        return False

    print(f"[+] Using {nessus_files[0]} as base structure")

    total_items = 0
    total_hosts = len(base_report.findall('ReportHost'))
    for report_host in base_report.findall('ReportHost'):
        total_items += len(report_host.findall('ReportItem'))

    for nessus_file in nessus_files[1:]:
        print(f"[*] Merging: {nessus_file}")
        tree = parse_nessus_file(nessus_file)
        if tree is None:
            print(f"[!] Skipping {nessus_file} due to parse error")
            continue
        report = tree.getroot().find('Report')
        if report is None:
            print(f"[!] No Report element in {nessus_file}, skipping")
            continue
        report_hosts = report.findall('ReportHost')
        for report_host in report_hosts:
            base_report.append(report_host)
            total_hosts += 1
            total_items += len(report_host.findall('ReportItem'))
        print(f"    Added {len(report_hosts)} host(s)")

    report_name = base_report.get('name', 'merged_scan')
    base_report.set('name', f"{report_name}_merged_{datetime.now().strftime('%Y%m%d_%H%M%S')}")

    try:
        ET.indent(base_tree, space="  ")
        base_tree.write(output_file, encoding='utf-8', xml_declaration=True)
        print(f"\n[+] Successfully merged {len(nessus_files)} file(s)")
        print(f"[+] Total hosts: {total_hosts}")
        print(f"[+] Total findings: {total_items}")
        print(f"[+] Output written to: {output_file}")
        return True
    except Exception as e:
        print(f"[!] Error writing output file: {e}")
        return False


def display_test_finding(organized_findings):
    print("\n" + "=" * 80)
    print("TEST MODE - Displaying First Parsed Finding")
    print("=" * 80 + "\n")

    for severity in SEVERITY_ORDER:
        if organized_findings[severity]:
            finding = organized_findings[severity][0]
            print(f"Severity: {finding.severity}")
            print(f"Finding: {finding.name}")
            print(f"Plugin ID: {finding.plugin_id}")
            if finding.cvss_score:
                print(f"CVSS Score: {finding.cvss_score}")
            print(f"\nAffected System(s):")
            for system in finding.systems():
                print(f"  - {system}")
            print(f"\nDescription:")
            print(finding.description[:500] + "..." if len(finding.description) > 500 else finding.description)
            print(f"\nRemediation:")
            print(finding.solution[:500] + "..." if len(finding.solution) > 500 else finding.solution)
            if finding.references:
                print(f"\nReferences:")
                for ref in finding.references[:10]:
                    print(f"  {ref}")
                if len(finding.references) > 10:
                    print(f"  ... and {len(finding.references) - 10} more")
            evidence = [out for out in finding.evidence_by_system.values() if out]
            if evidence:
                print(f"\nEvidence (first occurrence):")
                sample = evidence[0]
                print(sample[:500] + "..." if len(sample) > 500 else sample)
            print("\n" + "=" * 80)
            return

    print("[!] No findings found to display")


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Calibri (Headings)'
        if level == 2:
            run.font.size = Pt(13)
        elif level == 3:
            run.font.size = Pt(11)
    return heading


def add_paragraph(doc, text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        para.paragraph_format.left_indent = Inches(0.5 * level)
    for run in para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    return para


def add_label(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    return para


def add_kv(doc, key, value):
    para = doc.add_paragraph()
    k = para.add_run(key)
    k.bold = True
    k.font.name = 'Arial'
    k.font.size = Pt(11)
    v = para.add_run(value)
    v.font.name = 'Arial'
    v.font.size = Pt(11)
    return para


def add_na(doc, text='N/A'):
    para = add_paragraph(doc, text)
    para.runs[0].italic = True
    para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    return para


def add_evidence_block(doc, text):
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'D9D9D9')
    para._p.get_or_add_pPr().append(shading)
    return para


def render_evidence(doc, finding):
    if finding.evidence_override:
        add_evidence_block(doc, finding.evidence_override)
        return

    for system in finding.systems():
        output = finding.evidence_by_system.get(system, "")
        if output:
            add_evidence_block(doc, output)
            return

    add_na(doc)


def generate_report(organized_findings, output_file):
    print(f"[*] Generating DOCX report...")
    try:
        doc = Document()

        for severity in SEVERITY_ORDER:
            findings = organized_findings.get(severity, [])
            if not findings:
                continue

            add_heading(doc, f'{severity} Severity Findings', level=2)

            for finding in findings:
                add_heading(doc, finding.name, level=3)

                add_kv(doc, 'Severity: ', finding.severity)

                label = exploit_label(finding)
                if label:
                    add_kv(doc, 'Exploit available: ', label)

                add_label(doc, 'Affected System(s):')
                systems = finding.systems()
                if systems:
                    for system in systems:
                        add_bullet(doc, system)
                else:
                    add_bullet(doc, 'Unknown')

                add_label(doc, 'Description:')
                if finding.description:
                    add_paragraph(doc, finding.description)
                else:
                    add_na(doc)

                add_label(doc, 'Remediation:')
                if finding.solution:
                    add_paragraph(doc, finding.solution)
                else:
                    add_na(doc)

                add_label(doc, 'References:')
                if finding.references:
                    if finding.is_merged:
                        cve_ids = [ref.replace('CVE: ', '') for ref in finding.references]
                        add_paragraph(doc, ", ".join(cve_ids))
                    else:
                        for ref in finding.references:
                            para = doc.add_paragraph(ref)
                            for run in para.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                else:
                    add_na(doc, 'None')

                add_label(doc, 'Evidence:')
                render_evidence(doc, finding)

                doc.add_paragraph()

        doc.save(output_file)
        print(f"[+] Report generated: {output_file}")
        return True
    except Exception as e:
        print(f"[!] Error generating report: {e}")
        return False


class _PromptTimeout(Exception):
    pass


def _prompt_timeout_handler(signum, frame):
    raise _PromptTimeout()


def scan_files_at(path):
    if os.path.isfile(path):
        return [path] if path.endswith('.nessus') else []
    if os.path.isdir(path):
        return find_nessus_files(path)
    return []


def prompt_for_scan_path(timeout=30):
    signal.signal(signal.SIGALRM, _prompt_timeout_handler)
    signal.alarm(timeout)
    try:
        return input("[?] Enter path to a .nessus file or directory: ").strip()
    except _PromptTimeout:
        print(f"\n[!] No response within {timeout} seconds. Exiting.")
        sys.exit(1)
    except (EOFError, KeyboardInterrupt):
        print("\n[!] No input received. Exiting.")
        sys.exit(1)
    finally:
        signal.alarm(0)


def resolve_nessus_files(initial_path):
    scan_path = initial_path if initial_path else os.getcwd()
    nessus_files = scan_files_at(scan_path)

    if not nessus_files:
        if initial_path:
            print(f"[!] No .nessus file found at: {scan_path}")
        else:
            print(f"[*] No .nessus file found in {scan_path}")
        while not nessus_files:
            scan_path = prompt_for_scan_path()
            if not scan_path:
                print("[!] No path provided. Exiting.")
                sys.exit(1)
            nessus_files = scan_files_at(scan_path)
            if not nessus_files:
                print(f"[!] No .nessus file found at: {scan_path}")

    if len(nessus_files) == 1:
        print(f"[*] Using file: {nessus_files[0]}")
    else:
        print(f"[+] Found {len(nessus_files)} .nessus file(s):")
        for idx, file in enumerate(nessus_files, 1):
            file_size = os.path.getsize(file) / 1024
            print(f"    {idx}. {os.path.basename(file)} ({file_size:.1f} KB)")

    return nessus_files


def main():
    parser = argparse.ArgumentParser(
        description='Nessus File Merger & Report Generator',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Auto-detect .nessus files in current directory and generate report
  python3 nessus_parser.py

  # Specify directory or file
  python3 nessus_parser.py /path/to/scans

  # Test mode - view parsed data
  python3 nessus_parser.py --test

  # Merge without report
  python3 nessus_parser.py --no-report
        """
    )

    parser.add_argument('path', nargs='?', default=None,
                        help='Directory containing .nessus files or single .nessus file (default: current directory)')
    parser.add_argument('--output', '-o', default='merged_scan.nessus',
                        help='Output filename for merged .nessus file (default: merged_scan.nessus)')
    parser.add_argument('--report', '-r', action='store_true',
                        help='Generate report only (skip merge even with multiple files)')
    parser.add_argument('--test', '-t', action='store_true',
                        help='Test mode: display first parsed finding without generating report')
    parser.add_argument('--no-report', action='store_true',
                        help='Skip report generation (merge only)')
    parser.add_argument('--report-output', default=None,
                        help='Output filename for DOCX report (default: auto-generated from scan name)')

    args = parser.parse_args()

    nessus_files = resolve_nessus_files(args.path)

    if args.report_output:
        report_output = args.report_output
    else:
        if len(nessus_files) == 1:
            base_name = Path(nessus_files[0]).stem
        else:
            base_name = f"nessus_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        report_output = f"{base_name}.docx"

    if not args.report and len(nessus_files) > 1:
        print("\n[*] Merging .nessus files...")
        if not merge_nessus_files(nessus_files, args.output):
            print("[!] Merge failed")
            sys.exit(1)
        output_size = os.path.getsize(args.output) / 1024
        print(f"[+] Merged file size: {output_size:.1f} KB")

    print("\n[*] Parsing findings...")
    organized_findings, raw_item_count, unique_plugins, dropped_404, adjusted_count = parse_findings(nessus_files)

    total_findings = sum(len(findings) for findings in organized_findings.values())
    if dropped_404:
        print(f"[+] Dropped {dropped_404} false-positive instance(s) returning HTTP 404")
    if adjusted_count:
        print(f"[+] Adjusted {adjusted_count} CVE-backed finding(s) down one level (no public exploit)")
    print(f"[+] {raw_item_count} non-informational items -> {unique_plugins} unique plugins -> {total_findings} findings after condensation")
    for severity in SEVERITY_ORDER:
        count = len(organized_findings[severity])
        if count > 0:
            print(f"    {severity}: {count}")

    if args.test:
        display_test_finding(organized_findings)
        sys.exit(0)

    if not args.no_report:
        print(f"\n[*] Generating report: {report_output}")
        if generate_report(organized_findings, report_output):
            report_size = os.path.getsize(report_output) / 1024
            print(f"[+] Report size: {report_size:.1f} KB")
        else:
            print("[!] Report generation failed")
            sys.exit(1)

    print("\n[+] Complete")


if __name__ == "__main__":
    main()
