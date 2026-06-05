#!/usr/bin/env python3
# =============================================================================
# VAPT Toolkit - Vulnerability Assessment and Penetration Testing Toolkit
# =============================================================================
#
# Author: Keith Pachulski
# Company: Red Cell Security, LLC
# Email: keith@redcellsecurity.org
# Website: www.redcellsecurity.org
#
# Copyright (c) 2026 Keith Pachulski. All rights reserved.
#
# License: This software is licensed under the MIT License.
#          You are free to use, modify, and distribute this software
#          in accordance with the terms of the license.
#
# Purpose: Unified report parser that ingests Nessus, Burp Suite, and OWASP ZAP
#          XML output and completed penetration test reports (DOCX) from the
#          working directory and generates a consolidated DOCX vulnerability
#          assessment report and a DefectDojo Generic Findings Import JSON
#          export. Supports Nessus file merging and produces output suitable
#          for inclusion in formal client reports.
#
# DISCLAIMER: This software is provided "as-is," without warranty of any kind,
#             express or implied, including but not limited to the warranties
#             of merchantability, fitness for a particular purpose, and non-infringement.
#             In no event shall the authors or copyright holders be liable for any claim,
#             damages, or other liability, whether in an action of contract, tort, or otherwise,
#             arising from, out of, or in connection with the software or the use or other dealings
#             in the software.
#
# NOTICE: This toolkit is intended for authorized security testing only.
#         Users are responsible for ensuring compliance with all applicable laws
#         and regulations. Unauthorized use of these tools may violate local,
#         state, federal, and international laws.
#
# =============================================================================

import os
import re
import sys
import json
import hashlib
import argparse
import xml.etree.ElementTree as ET
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field
from urllib.parse import urlparse

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph


# Severity ordering used in output. High/Medium/Low only - Informational is
# dropped at parse time, Critical collapses into High.
SEVERITY_ORDER = ['High', 'Medium', 'Low']

SEVERITY_COLORS = {
    'High': RGBColor(255, 0, 0),
    'Medium': RGBColor(255, 140, 0),
    'Low': RGBColor(65, 105, 225),
}

# Per-scanner severity normalization. Anything not in these maps (Informational,
# Information, None, empty) is dropped at parse time.
NESSUS_SEVERITY_MAP = {
    'Critical': 'High',
    'High': 'High',
    'Medium': 'Medium',
    'Low': 'Low',
}

BURP_SEVERITY_MAP = {
    'High': 'High',
    'Medium': 'Medium',
    'Low': 'Low',
}

ZAP_SEVERITY_MAP = {
    'High': 'High',
    'Medium': 'Medium',
    'Low': 'Low',
}

# DefectDojo Generic Findings Import default test type name
DEFAULT_TEST_TYPE_NAME = 'VAPT Assessment'

# Labels that mark field boundaries inside a finding in a formal report. Order
# matters for the state machine: once a label is hit, all subsequent content
# belongs to it until the next label, the next H3 (new finding), or the next
# H1/H2 (section end).
REPORT_FIELD_LABELS = ('description', 'recommendation', 'references', 'evidence')


@dataclass
class Finding:
    """Normalized finding representation used across all scanner types."""
    scanner_type: str            # 'nessus' | 'burp' | 'zap' | 'manual'
    scanner_id: str              # plugin_id for nessus, finding name for others
    title: str
    severity: str                # 'High' | 'Medium' | 'Low'
    affected_systems: list = field(default_factory=list)
    description: str = ''
    recommendation: str = ''
    references: list = field(default_factory=list)
    evidence_blocks: list = field(default_factory=list)  # [{'label': str, 'content': str}]

    # Parsed but not rendered in DOCX. Retained for DefectDojo JSON export.
    cvss_score: str = ''
    cwe: str = ''
    cve: str = ''
    plugin_id: str = ''


# ---------------------------------------------------------------------------
# Format detection and file discovery
# ---------------------------------------------------------------------------

def detect_format(file_path):
    """
    Identify scanner output type. .docx files are checked by extension since
    they're ZIP archives, not text-sniffable. XML files are sniffed.
    Returns 'nessus' | 'burp' | 'zap' | 'report' | None.
    """
    if file_path.suffix.lower() == '.docx':
        return 'report'

    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read(2048)
    except (IOError, OSError):
        return None

    if '<NessusClientData_v2' in content:
        return 'nessus'
    if '<OWASPZAPReport' in content or 'programName="ZAP"' in content:
        return 'zap'
    if '<!DOCTYPE issues' in content or 'burpVersion' in content:
        return 'burp'
    return None


def find_scanner_files(directory, include_reports=False):
    """
    Scan the working directory (top level only, no recursion) for scanner output
    and optionally completed penetration test reports.

    Returns a list of (Path, scanner_type) tuples. DOCX report files are only
    included when include_reports=True, since DOCX files in a working directory
    may be unrelated to the engagement (SOWs, notes, prior reports).
    """
    results = []
    for entry in sorted(Path(directory).iterdir()):
        if not entry.is_file():
            continue
        if entry.suffix.lower() not in ['.nessus', '.xml', '.docx']:
            continue
        scanner_type = detect_format(entry)
        if not scanner_type:
            continue
        if scanner_type == 'report' and not include_reports:
            continue
        results.append((entry, scanner_type))
    return results


# ---------------------------------------------------------------------------
# Shared parsing helpers
# ---------------------------------------------------------------------------

def _extract_text(element, tag):
    """Safely extract text content from a named child element."""
    child = element.find(tag)
    return child.text if child is not None and child.text else ''


def _strip_html(text):
    """Remove HTML tags from a string."""
    if not text:
        return ''
    return re.sub(r'<[^>]+>', '', text).strip()


def _extract_base_url(uri):
    """Reduce a URI to scheme://host[:port], preserving non-standard ports only."""
    if not uri:
        return ''
    parsed = urlparse(uri)
    if not parsed.hostname:
        return uri
    base = f"{parsed.scheme}://{parsed.hostname}"
    if parsed.port:
        if (parsed.scheme == 'https' and parsed.port != 443) or \
           (parsed.scheme == 'http' and parsed.port != 80):
            base += f":{parsed.port}"
    return base


# ---------------------------------------------------------------------------
# Nessus adapter
# ---------------------------------------------------------------------------

def parse_nessus(file_path):
    """Parse a .nessus XML file and return a list of Finding objects."""
    try:
        tree = ET.parse(file_path)
    except ET.ParseError as e:
        print(f"[!] Parse error in {file_path}: {e}")
        return []

    root = tree.getroot()
    report = root.find('Report')
    if report is None:
        return []

    findings = []

    for report_host in report.findall('ReportHost'):
        host_name = report_host.get('name', 'Unknown')

        for report_item in report_host.findall('ReportItem'):
            plugin_id = report_item.get('pluginID', '')
            raw_severity = _extract_text(report_item, 'risk_factor')
            severity = NESSUS_SEVERITY_MAP.get(raw_severity)
            if not severity:
                continue

            # Build affected system string: host:port (protocol/svc)
            port = report_item.get('port', '')
            protocol = report_item.get('protocol', '')
            svc_name = report_item.get('svc_name', '')
            system_info = host_name
            if port and port != '0':
                system_info += f":{port}"
            if protocol:
                system_info += f" ({protocol}"
                if svc_name:
                    system_info += f"/{svc_name}"
                system_info += ")"

            # References: CVE, xref, see_also
            refs = []
            cve_value = ''
            for cve in report_item.findall('cve'):
                if cve.text:
                    refs.append(f"CVE: {cve.text}")
                    if not cve_value:
                        cve_value = cve.text
            for xref in report_item.findall('xref'):
                if xref.text:
                    refs.append(xref.text)
            for see_also in report_item.findall('see_also'):
                if see_also.text:
                    refs.append(see_also.text)

            plugin_output = _extract_text(report_item, 'plugin_output')
            evidence_blocks = []
            if plugin_output:
                evidence_blocks.append({'label': '', 'content': plugin_output})

            finding = Finding(
                scanner_type='nessus',
                scanner_id=plugin_id,
                title=report_item.get('pluginName', 'Unknown'),
                severity=severity,
                affected_systems=[system_info],
                description=_extract_text(report_item, 'description'),
                recommendation=_extract_text(report_item, 'solution'),
                references=refs,
                evidence_blocks=evidence_blocks,
                cvss_score=_extract_text(report_item, 'cvss_base_score'),
                cve=cve_value,
                plugin_id=plugin_id,
            )
            findings.append(finding)

    return findings


# ---------------------------------------------------------------------------
# Burp Suite adapter
# ---------------------------------------------------------------------------

def parse_burp(file_path):
    """Parse a Burp Suite XML report and return a list of Finding objects."""
    try:
        tree = ET.parse(file_path)
    except ET.ParseError as e:
        print(f"[!] Parse error in {file_path}: {e}")
        return []

    root = tree.getroot()
    findings = []

    for issue in root.findall('.//issue'):
        raw_severity = (issue.findtext('severity') or '').strip()
        severity = BURP_SEVERITY_MAP.get(raw_severity)
        if not severity:
            continue

        name = issue.findtext('name', 'Unknown')
        host = issue.findtext('host', '')
        path = issue.findtext('path', '')

        affected = []
        if host:
            full_url = f"{host}{path}" if path else host
            affected.append(_extract_base_url(full_url) or full_url)

        description = (issue.findtext('issueBackground', '') or
                       issue.findtext('issueDetail', ''))
        recommendation = (issue.findtext('remediationBackground', '') or
                          issue.findtext('remediationDetail', ''))

        # References can appear as a single <references> blob or individual <reference> elements.
        refs = []
        ref_text = issue.findtext('references', '')
        if ref_text:
            for line in _strip_html(ref_text).split('\n'):
                line = line.strip()
                if line:
                    refs.append(line)
        for ref_elem in issue.findall('.//reference'):
            if ref_elem.text and ref_elem.text.strip():
                refs.append(ref_elem.text.strip())

        # Capture CWE number if present in vulnerability classifications.
        cwe = ''
        vc_text = issue.findtext('vulnerabilityClassifications', '') or ''
        cwe_match = re.search(r'CWE-(\d+)', vc_text)
        if cwe_match:
            cwe = cwe_match.group(1)

        # Request/response from the first requestresponse element.
        evidence_blocks = []
        rr = issue.find('.//requestresponse')
        if rr is not None:
            request_elem = rr.find('request')
            response_elem = rr.find('response')
            if request_elem is not None and request_elem.text:
                evidence_blocks.append({
                    'label': 'Request',
                    'content': request_elem.text,
                })
            if response_elem is not None and response_elem.text:
                body = response_elem.text
                if len(body) > 1000:
                    body = body[:1000] + '\n... (truncated)'
                evidence_blocks.append({
                    'label': 'Response',
                    'content': body,
                })

        finding = Finding(
            scanner_type='burp',
            scanner_id=name,
            title=name,
            severity=severity,
            affected_systems=affected,
            description=_strip_html(description),
            recommendation=_strip_html(recommendation),
            references=refs,
            evidence_blocks=evidence_blocks,
            cwe=cwe,
        )
        findings.append(finding)

    return findings


# ---------------------------------------------------------------------------
# OWASP ZAP adapter
# ---------------------------------------------------------------------------

def parse_zap(file_path):
    """Parse an OWASP ZAP XML report and return a list of Finding objects."""
    try:
        tree = ET.parse(file_path)
    except ET.ParseError as e:
        print(f"[!] Parse error in {file_path}: {e}")
        return []

    root = tree.getroot()
    findings = []

    for site in root.findall('.//site'):
        for alert in site.findall('.//alertitem'):
            riskdesc = alert.findtext('riskdesc', '') or ''
            risk_raw = riskdesc.split()[0] if riskdesc else ''
            severity = ZAP_SEVERITY_MAP.get(risk_raw)
            if not severity:
                continue

            name = alert.findtext('name', 'Unknown')
            description = _strip_html(alert.findtext('desc', ''))
            recommendation = _strip_html(alert.findtext('solution', ''))

            refs = []
            ref_text = alert.findtext('reference', '')
            if ref_text:
                for line in _strip_html(ref_text).split('\n'):
                    line = line.strip()
                    if line:
                        refs.append(line)

            cwe = (alert.findtext('cweid', '') or '').strip()

            affected = []
            request_combined = ''
            response_combined = ''
            for instance in alert.findall('.//instance'):
                uri = instance.findtext('uri', '')
                if uri:
                    base = _extract_base_url(uri)
                    if base and base not in affected:
                        affected.append(base)
                # Capture first instance evidence only.
                if not request_combined and not response_combined:
                    req_h = instance.findtext('requestheader', '') or ''
                    req_b = instance.findtext('requestbody', '') or ''
                    res_h = instance.findtext('responseheader', '') or ''
                    res_b = instance.findtext('responsebody', '') or ''
                    if req_h or req_b:
                        request_combined = req_h
                        if req_b:
                            request_combined += ('\n\n' + req_b) if request_combined else req_b
                    if res_h or res_b:
                        response_combined = res_h
                        if res_b:
                            if len(res_b) > 1000:
                                res_b = res_b[:1000] + '\n... (truncated)'
                            response_combined += ('\n\n' + res_b) if response_combined else res_b

            evidence_blocks = []
            if request_combined:
                evidence_blocks.append({'label': 'Request', 'content': request_combined})
            if response_combined:
                evidence_blocks.append({'label': 'Response', 'content': response_combined})

            finding = Finding(
                scanner_type='zap',
                scanner_id=name,
                title=name,
                severity=severity,
                affected_systems=affected,
                description=description,
                recommendation=recommendation,
                references=refs,
                evidence_blocks=evidence_blocks,
                cwe=cwe,
            )
            findings.append(finding)

    return findings


# ---------------------------------------------------------------------------
# Manual penetration test report (DOCX) adapter
# ---------------------------------------------------------------------------

def _iter_block_items(parent):
    """
    Yield paragraphs and tables in document order from a python-docx parent
    (Document or _Cell). python-docx exposes paragraphs and tables as separate
    lists, which loses ordering; this walks the underlying XML to preserve it.
    """
    from docx.document import Document as _Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _heading_level(paragraph):
    """Return heading level as int (1, 2, 3, ...) or None if not a heading."""
    style_name = paragraph.style.name if paragraph.style else ''
    if not style_name.startswith('Heading '):
        return None
    try:
        return int(style_name.replace('Heading ', '').strip())
    except (ValueError, AttributeError):
        return None


def _paragraph_text(paragraph):
    """
    Extract paragraph text while stripping images and inline drawings. Images
    are dropped silently. Other run content is preserved.
    """
    parts = []
    for run in paragraph.runs:
        if run.element.findall('.//' + qn('w:drawing')):
            if not run.text:
                continue
        parts.append(run.text or '')
    return ''.join(parts)


def _is_bold_label(paragraph, label):
    """
    Check if a paragraph is a bold label matching the given text (case-insensitive,
    colon-tolerant).
    """
    text = _paragraph_text(paragraph).strip().rstrip(':').strip().lower()
    if text != label.lower():
        return False
    return any(run.bold for run in paragraph.runs)


def _cell_text_lines(cell):
    """Return a list of non-empty text lines from a table cell, one per paragraph."""
    lines = []
    for para in cell.paragraphs:
        text = _paragraph_text(para).strip()
        if text:
            lines.append(text)
    return lines


def _split_affected_systems(cell):
    """
    Split a metadata-table affected-systems cell into individual entries.
    Each paragraph in the cell is treated as one system entry, preserving any
    inline annotation (e.g., 'https://example.com/path  (parameter: filter)').
    """
    systems = []
    for line in _cell_text_lines(cell):
        line = line.strip()
        if line and line not in systems:
            systems.append(line)
    return systems


def _parse_metadata_table(table):
    """
    Parse the per-finding metadata table.
    Returns (severity_raw, affected_systems_list) or (None, []).
    """
    severity = None
    affected = []

    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = _paragraph_text(row.cells[0].paragraphs[0]).strip().rstrip(':').strip().lower()

        if label == 'severity':
            sev_text = _paragraph_text(row.cells[1].paragraphs[0]).strip()
            severity = sev_text
        elif label in ('affected system(s)', 'affected systems', 'affected system'):
            affected = _split_affected_systems(row.cells[1])

    return severity, affected


def _normalize_report_severity(raw):
    """Normalize a free-text severity string to High/Medium/Low or None."""
    if not raw:
        return None
    raw = raw.strip().lower()
    if raw in ('critical', 'high'):
        return 'High'
    if raw == 'medium':
        return 'Medium'
    if raw == 'low':
        return 'Low'
    return None


def parse_report(file_path):
    """
    Parse a finished penetration test report (DOCX) and return Finding objects.

    Expects the formal report structure:
        H1 'Detailed Findings'
            H2 '<Severity> Severity Findings'
                H3 '<Finding Title>'
                    Metadata table (Severity, Affected System(s))
                    Bold 'Description' -> prose
                    Bold 'Recommendation' -> prose
                    Bold 'References' -> bullets
                    Bold 'Evidence' -> mixed content
        H1 '<next section>' (terminates findings)
    """
    try:
        doc = Document(str(file_path))
    except Exception as e:
        print(f"[!] Failed to open {file_path}: {e}")
        return []

    findings = []
    blocks = list(_iter_block_items(doc))

    # Locate the start of the findings section
    start_idx = None
    for idx, block in enumerate(blocks):
        if not isinstance(block, DocxParagraph):
            continue
        if _heading_level(block) == 1:
            heading_text = _paragraph_text(block).strip().lower()
            if 'detailed findings' in heading_text:
                start_idx = idx + 1
                break

    if start_idx is None:
        print(f"[!] No 'Detailed Findings' section found in {file_path.name}")
        return []

    current_finding = None
    current_field = None
    field_buffers = {label: [] for label in REPORT_FIELD_LABELS}
    field_references = []

    def finalize_current():
        if not current_finding:
            return
        description = '\n\n'.join(field_buffers['description']).strip()
        recommendation = '\n\n'.join(field_buffers['recommendation']).strip()
        evidence_text = '\n'.join(field_buffers['evidence']).rstrip()

        current_finding.description = description
        current_finding.recommendation = recommendation
        current_finding.references = list(field_references)
        if evidence_text:
            current_finding.evidence_blocks = [{'label': '', 'content': evidence_text}]
        findings.append(current_finding)

    for block in blocks[start_idx:]:
        # Table handling: a table immediately after an H3 is the metadata table
        if isinstance(block, DocxTable):
            if current_finding and current_field is None:
                severity_raw, affected = _parse_metadata_table(block)
                severity = _normalize_report_severity(severity_raw)
                if severity:
                    current_finding.severity = severity
                if affected:
                    current_finding.affected_systems = affected
            continue

        if not isinstance(block, DocxParagraph):
            continue

        level = _heading_level(block)
        text = _paragraph_text(block).strip()

        # H1 terminates the findings section
        if level == 1:
            finalize_current()
            current_finding = None
            current_field = None
            field_buffers = {label: [] for label in REPORT_FIELD_LABELS}
            field_references = []
            break

        # H2 is informational only; severity comes from the metadata table
        if level == 2:
            continue

        # H3 starts a new finding
        if level == 3:
            finalize_current()
            current_finding = Finding(
                scanner_type='manual',
                scanner_id=text,
                title=text,
                severity='',
            )
            current_field = None
            field_buffers = {label: [] for label in REPORT_FIELD_LABELS}
            field_references = []
            continue

        if current_finding is None:
            continue

        label_match = None
        for label in REPORT_FIELD_LABELS:
            if _is_bold_label(block, label):
                label_match = label
                break

        if label_match:
            current_field = label_match
            continue

        if current_field is None:
            continue

        if current_field == 'references':
            if text:
                field_references.append(text)
        else:
            # Evidence preserves blank lines for HTTP request/response readability
            if text or current_field == 'evidence':
                field_buffers[current_field].append(text)

    # Flush trailing finding if document ended without another H1
    finalize_current()

    # Drop findings whose severity didn't normalize
    valid_findings = [f for f in findings if f.severity in SEVERITY_ORDER]
    dropped = len(findings) - len(valid_findings)
    if dropped:
        print(f"    Skipped {dropped} finding(s) with non-mapped severity")

    return valid_findings


# ---------------------------------------------------------------------------
# Nessus XML merge
# ---------------------------------------------------------------------------

def merge_nessus_files(nessus_files, output_file):
    """
    Merge multiple .nessus files into a single output file.
    Uses the Policy from the first file and combines all ReportHost elements.
    """
    if not nessus_files:
        print("[!] No .nessus files to merge")
        return False

    try:
        base_tree = ET.parse(nessus_files[0])
    except ET.ParseError as e:
        print(f"[!] Failed to parse base file {nessus_files[0]}: {e}")
        return False

    base_root = base_tree.getroot()
    base_report = base_root.find('Report')
    if base_report is None:
        print(f"[!] No Report element in base file {nessus_files[0]}")
        return False

    print(f"[+] Using {nessus_files[0]} as base structure")

    total_hosts = len(base_report.findall('ReportHost'))
    total_items = sum(len(h.findall('ReportItem')) for h in base_report.findall('ReportHost'))

    for nessus_file in nessus_files[1:]:
        print(f"[*] Merging: {nessus_file}")
        try:
            tree = ET.parse(nessus_file)
        except ET.ParseError as e:
            print(f"[!] Skipping {nessus_file} due to parse error: {e}")
            continue

        report = tree.getroot().find('Report')
        if report is None:
            print(f"[!] No Report element in {nessus_file}, skipping")
            continue

        report_hosts = report.findall('ReportHost')
        for report_host in report_hosts:
            base_report.append(report_host)
            total_hosts += 1
            total_items += len(report_host.findall('ReportItem'))

        print(f"    Added {len(report_hosts)} host(s)")

    report_name = base_report.get('name', 'merged_scan')
    base_report.set('name', f"{report_name}_merged_{datetime.now().strftime('%Y%m%d_%H%M%S')}")

    try:
        ET.indent(base_tree, space="  ")
        base_tree.write(output_file, encoding='utf-8', xml_declaration=True)
        print(f"[+] Merged {len(nessus_files)} file(s) -> {output_file}")
        print(f"[+] Total hosts: {total_hosts}, total findings: {total_items}")
        return True
    except (IOError, OSError) as e:
        print(f"[!] Error writing merged file: {e}")
        return False


# ---------------------------------------------------------------------------
# Aggregation and dispatch
# ---------------------------------------------------------------------------

def aggregate_findings(findings):
    """
    Deduplicate findings within each scanner type and merge affected systems
    across duplicates. Cross-scanner duplicates are preserved as separate findings.
    """
    deduped = {}

    for finding in findings:
        key = (finding.scanner_type, finding.scanner_id)
        if key not in deduped:
            deduped[key] = finding
            continue

        existing = deduped[key]
        for system in finding.affected_systems:
            if system not in existing.affected_systems:
                existing.affected_systems.append(system)

    return list(deduped.values())


def bucket_by_severity(findings):
    """
    Organize findings into severity buckets, sorted alphabetically by title
    within each bucket. Returns {'High': [...], 'Medium': [...], 'Low': [...]}.
    """
    buckets = {sev: [] for sev in SEVERITY_ORDER}
    for finding in findings:
        if finding.severity in buckets:
            buckets[finding.severity].append(finding)
    for sev in SEVERITY_ORDER:
        buckets[sev].sort(key=lambda f: f.title.lower())
    return buckets


def parse_all_files(scanner_files):
    """
    Run the appropriate adapter against each file and return a flat list of
    Finding objects from all sources combined.
    """
    all_findings = []
    adapter_map = {
        'nessus': parse_nessus,
        'burp': parse_burp,
        'zap': parse_zap,
        'report': parse_report,
    }

    for file_path, scanner_type in scanner_files:
        adapter = adapter_map.get(scanner_type)
        if not adapter:
            continue
        print(f"[*] Parsing {scanner_type.upper()}: {file_path.name}")
        file_findings = adapter(file_path)
        print(f"    Extracted {len(file_findings)} finding(s)")
        all_findings.extend(file_findings)

    return all_findings


# ---------------------------------------------------------------------------
# DOCX report generator
# ---------------------------------------------------------------------------

def _set_cell_shading(paragraph, fill_color):
    """Apply background shading to a paragraph (used for evidence blocks)."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    paragraph._p.get_or_add_pPr().append(shading)


def _add_na(doc, text='N/A'):
    """Add an italic grey 'N/A' or 'None' paragraph for empty fields."""
    para = doc.add_paragraph(text)
    para.runs[0].italic = True
    para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    return para


def _add_label(doc, label_text):
    """Add a bold label paragraph (e.g., 'Description:')."""
    para = doc.add_paragraph()
    run = para.add_run(label_text)
    run.bold = True
    return para


def _add_evidence_block(doc, content):
    """Add a monospace, grey-shaded evidence paragraph."""
    para = doc.add_paragraph(content)
    for run in para.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
    _set_cell_shading(para, 'D9D9D9')
    return para


def generate_report(buckets, output_file):
    """Generate the DOCX report from severity-bucketed findings."""
    print("[*] Generating DOCX report...")

    try:
        doc = Document()

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Title
        title = doc.add_heading('Vulnerability Assessment Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Timestamp
        timestamp = doc.add_paragraph()
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = timestamp.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(102, 102, 102)

        # Executive Summary
        doc.add_heading('Executive Summary', level=1)

        total = sum(len(buckets[sev]) for sev in SEVERITY_ORDER)
        doc.add_paragraph(
            f'This report contains {total} unique security findings organized by severity:'
        )

        for severity in SEVERITY_ORDER:
            count = len(buckets[severity])
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(f'{severity}: {count}')
            run.bold = True
            run.font.color.rgb = SEVERITY_COLORS[severity]

        # Findings
        doc.add_heading('Findings', level=1)

        for severity in SEVERITY_ORDER:
            findings = buckets[severity]
            if not findings:
                continue

            sev_heading = doc.add_heading(severity, level=2)
            sev_heading.runs[0].font.color.rgb = SEVERITY_COLORS[severity]

            for finding in findings:
                doc.add_heading(finding.title, level=3)

                # Affected System(s)
                _add_label(doc, 'Affected System(s):')
                if finding.affected_systems:
                    for system in finding.affected_systems:
                        doc.add_paragraph(system, style='List Bullet')
                else:
                    _add_na(doc, 'None')

                # Description
                _add_label(doc, 'Description:')
                if finding.description:
                    doc.add_paragraph(finding.description)
                else:
                    _add_na(doc)

                # Recommendation
                _add_label(doc, 'Recommendation:')
                if finding.recommendation:
                    doc.add_paragraph(finding.recommendation)
                else:
                    _add_na(doc)

                # References
                _add_label(doc, 'References:')
                if finding.references:
                    for ref in finding.references:
                        doc.add_paragraph(ref)
                else:
                    _add_na(doc, 'None')

                # Evidence
                _add_label(doc, 'Evidence:')
                if finding.evidence_blocks:
                    for block in finding.evidence_blocks:
                        if block['label']:
                            sub = doc.add_paragraph()
                            sub.add_run(f"{block['label']}:").bold = True
                        _add_evidence_block(doc, block['content'])
                else:
                    _add_na(doc)

                # Spacer between findings
                doc.add_paragraph()

        doc.save(output_file)
        print(f"[+] Report written: {output_file}")
        return True

    except Exception as e:
        print(f"[!] Error generating report: {e}")
        return False


# ---------------------------------------------------------------------------
# DefectDojo JSON exporter
# ---------------------------------------------------------------------------

def _generate_unique_id(finding):
    """
    Generate a stable unique_id_from_tool for DefectDojo remediation tracking.

    Nessus uses plugin_id directly. Burp/ZAP/manual hash the finding title since
    none expose a stable vulnerability ID in their source.
    """
    if finding.scanner_type == 'nessus':
        return f"nessus-{finding.plugin_id}"

    title_hash = hashlib.sha256(finding.title.encode('utf-8')).hexdigest()[:16]
    return f"{finding.scanner_type}-{title_hash}"


def _normalize_endpoint(system):
    """
    Normalize an affected system string for the DefectDojo endpoints array.

    Nessus format 'host:port (protocol/svc)' is stripped to 'host:port'.
    URLs are passed through unchanged.
    """
    if not system:
        return ''
    paren_idx = system.find(' (')
    if paren_idx > 0:
        return system[:paren_idx].strip()
    return system.strip()


def _build_finding_json(finding):
    """Convert a Finding object into a DefectDojo Generic Findings Import dict."""
    entry = {
        'title': finding.title,
        'description': finding.description or 'No description provided.',
        'severity': finding.severity,
        'date': datetime.now().strftime('%Y-%m-%d'),
        'unique_id_from_tool': _generate_unique_id(finding),
        'active': True,
        'verified': True,
        'static_finding': False,
        'dynamic_finding': True,
    }

    if finding.recommendation:
        entry['mitigation'] = finding.recommendation

    if finding.references:
        entry['references'] = '\n'.join(finding.references)

    if finding.plugin_id:
        entry['vuln_id_from_tool'] = finding.plugin_id

    if finding.cwe:
        try:
            entry['cwe'] = int(finding.cwe)
        except (ValueError, TypeError):
            pass

    if finding.cve:
        entry['cve'] = finding.cve

    if finding.cvss_score:
        try:
            entry['cvssv3_score'] = float(finding.cvss_score)
        except (ValueError, TypeError):
            if finding.cvss_score.startswith('CVSS:'):
                entry['cvssv3'] = finding.cvss_score

    # Endpoints from affected_systems
    endpoints = []
    for system in finding.affected_systems:
        normalized = _normalize_endpoint(system)
        if normalized and normalized not in endpoints:
            endpoints.append(normalized)
    if endpoints:
        entry['endpoints'] = endpoints

    # Evidence blocks into steps_to_reproduce
    if finding.evidence_blocks:
        parts = []
        for block in finding.evidence_blocks:
            if block['label']:
                parts.append(f"{block['label']}:\n{block['content']}")
            else:
                parts.append(block['content'])
        entry['steps_to_reproduce'] = '\n\n'.join(parts)

    # Scanner type as a tag for filtering inside DefectDojo
    entry['tags'] = [finding.scanner_type]

    return entry


def export_defectdojo_json(buckets, output_file, test_type_name):
    """Export deduplicated findings as DefectDojo Generic Findings Import JSON."""
    print(f"[*] Exporting DefectDojo JSON (test type: {test_type_name})...")

    findings_list = []
    for severity in SEVERITY_ORDER:
        for finding in buckets[severity]:
            findings_list.append(_build_finding_json(finding))

    payload = {
        'name': test_type_name,
        'findings': findings_list,
    }

    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(payload, f, indent=2, ensure_ascii=False)
        print(f"[+] DefectDojo JSON written: {output_file}")
        print(f"[+] Exported {len(findings_list)} finding(s)")
        return True
    except (IOError, OSError) as e:
        print(f"[!] Error writing JSON: {e}")
        return False


# ---------------------------------------------------------------------------
# Test-mode display
# ---------------------------------------------------------------------------

def display_test_finding(buckets):
    """Display the first finding from the highest-severity non-empty bucket."""
    print("\n" + "=" * 80)
    print("TEST MODE - First Parsed Finding")
    print("=" * 80 + "\n")

    for severity in SEVERITY_ORDER:
        if not buckets[severity]:
            continue

        finding = buckets[severity][0]
        print(f"Scanner: {finding.scanner_type}")
        print(f"Severity: {finding.severity}")
        print(f"Title: {finding.title}")
        print(f"Scanner ID: {finding.scanner_id}")

        print(f"\nAffected System(s) ({len(finding.affected_systems)}):")
        for system in finding.affected_systems[:10]:
            print(f"  - {system}")
        if len(finding.affected_systems) > 10:
            print(f"  ... and {len(finding.affected_systems) - 10} more")

        print(f"\nDescription:")
        desc = finding.description
        print(desc[:500] + ('...' if len(desc) > 500 else ''))

        print(f"\nRecommendation:")
        rec = finding.recommendation
        print(rec[:500] + ('...' if len(rec) > 500 else ''))

        if finding.references:
            print(f"\nReferences ({len(finding.references)}):")
            for ref in finding.references[:10]:
                print(f"  {ref}")
            if len(finding.references) > 10:
                print(f"  ... and {len(finding.references) - 10} more")

        if finding.evidence_blocks:
            print(f"\nEvidence Blocks: {len(finding.evidence_blocks)}")
            for block in finding.evidence_blocks:
                label = block['label'] or '(unlabeled)'
                content = block['content']
                preview = content[:200] + ('...' if len(content) > 200 else '')
                print(f"  [{label}] {preview}")

        print("\n" + "=" * 80)
        return

    print("[!] No findings to display")


# ---------------------------------------------------------------------------
# CLI / main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='VAPT Report Parser - Nessus / Burp / ZAP XML and manual report DOCX to DOCX and DefectDojo JSON',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Scan current working directory and generate DOCX report from scanner output
  python3 vapt_report_parser.py

  # Generate DOCX and DefectDojo JSON from scanner output
  python3 vapt_report_parser.py --json

  # Parse a completed penetration test report and export to DefectDojo
  python3 vapt_report_parser.py --include-reports --json --client "Acme Corp"

  # Test mode - display first finding without writing files
  python3 vapt_report_parser.py --test
        """
    )

    parser.add_argument('-o', '--output', default=None,
                        help='DOCX output path (default: vapt_report_<timestamp>.docx in working dir)')
    parser.add_argument('--merge-output', default='merged_scan.nessus',
                        help='Nessus merge output filename (default: merged_scan.nessus)')
    parser.add_argument('--no-merge', action='store_true',
                        help='Skip Nessus XML merge step')
    parser.add_argument('--no-report', action='store_true',
                        help='Skip DOCX report generation')
    parser.add_argument('--json', nargs='?', const='__auto__', default=None,
                        help='Export DefectDojo Generic Findings Import JSON. '
                             'Optional path argument (default: vapt_findings_<timestamp>.json)')
    parser.add_argument('--client', default=None,
                        help='Client name used as DefectDojo test type. '
                             'Default test type: "VAPT Assessment"')
    parser.add_argument('--include-reports', action='store_true',
                        help='Include completed penetration test reports (.docx) as input. '
                             'Off by default to avoid parsing unrelated DOCX files in the working directory.')
    parser.add_argument('--test', action='store_true',
                        help='Display first parsed finding and exit (no files written)')

    args = parser.parse_args()

    working_dir = Path.cwd()
    print(f"[*] Working directory: {working_dir}")

    scanner_files = find_scanner_files(working_dir, include_reports=args.include_reports)

    if not scanner_files:
        print("[!] No supported scanner files found in working directory")
        print("    Expected: .nessus (Nessus), .xml (Burp/ZAP)")
        if not args.include_reports:
            print("    Pass --include-reports to also parse completed report DOCX files")
        sys.exit(1)

    by_type = {'nessus': [], 'burp': [], 'zap': [], 'report': []}
    for file_path, scanner_type in scanner_files:
        by_type[scanner_type].append(file_path)

    print(f"[+] Found {len(scanner_files)} scanner file(s):")
    for scanner_type in ['nessus', 'burp', 'zap', 'report']:
        files = by_type[scanner_type]
        if files:
            print(f"    {scanner_type.upper()}: {len(files)}")
            for fp in files:
                size_kb = os.path.getsize(fp) / 1024
                print(f"      - {fp.name} ({size_kb:.1f} KB)")

    # Merge Nessus files if multiple present and not disabled
    if not args.no_merge and not args.test and len(by_type['nessus']) > 1:
        print(f"\n[*] Merging {len(by_type['nessus'])} Nessus file(s)...")
        merge_output_path = working_dir / args.merge_output
        if not merge_nessus_files([str(p) for p in by_type['nessus']], str(merge_output_path)):
            print("[!] Nessus merge failed (continuing with individual files for parsing)")

    # Parse all files
    print("\n[*] Parsing findings...")
    raw_findings = parse_all_files(scanner_files)

    if not raw_findings:
        print("[!] No findings extracted from any file")
        sys.exit(1)

    # Dedup and bucket
    deduped = aggregate_findings(raw_findings)
    buckets = bucket_by_severity(deduped)

    # Stats
    total = sum(len(buckets[sev]) for sev in SEVERITY_ORDER)
    print(f"\n[+] {len(raw_findings)} raw findings -> {total} unique after dedup")
    for severity in SEVERITY_ORDER:
        count = len(buckets[severity])
        if count > 0:
            print(f"    {severity}: {count}")

    # Test mode short-circuit
    if args.test:
        display_test_finding(buckets)
        sys.exit(0)

    # DOCX report generation
    # Auto-skip if the only inputs were completed reports - regenerating a DOCX
    # from a report's own findings produces a less complete copy of the input.
    only_reports = all(stype == 'report' for _, stype in scanner_files)

    if args.no_report:
        print("\n[*] DOCX report generation skipped (--no-report)")
    elif only_reports:
        print("\n[*] DOCX report generation skipped (input is report-only; use --json for DefectDojo export)")
    else:
        if args.output:
            output_path = Path(args.output)
            if not output_path.is_absolute():
                output_path = working_dir / output_path
        else:
            output_path = working_dir / f"vapt_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        print(f"\n[*] Generating report: {output_path}")
        if generate_report(buckets, str(output_path)):
            size_kb = os.path.getsize(output_path) / 1024
            print(f"[+] Report size: {size_kb:.1f} KB")
        else:
            print("[!] Report generation failed")
            sys.exit(1)

    # DefectDojo JSON export
    if args.json is not None:
        if args.json == '__auto__':
            json_path = working_dir / f"vapt_findings_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        else:
            json_path = Path(args.json)
            if not json_path.is_absolute():
                json_path = working_dir / json_path

        test_type_name = args.client if args.client else DEFAULT_TEST_TYPE_NAME
        print(f"\n[*] Exporting DefectDojo JSON: {json_path}")
        if not export_defectdojo_json(buckets, str(json_path), test_type_name):
            print("[!] JSON export failed")
            sys.exit(1)

    print("\n[+] Complete")


if __name__ == "__main__":
    main()
