#!/usr/bin/env python3
# =============================================================================
# VAPT Toolkit - Vulnerability Assessment and Penetration Testing Toolkit
# =============================================================================
#
# Author: Keith Pachulski
# Company: Red Cell Security, LLC
# Email: keith@redcellsecurity.org
# Website: www.redcellsecurity.org
#
# Copyright (c) 2025 Keith Pachulski. All rights reserved.
#
# License: This software is licensed under the MIT License.
#          You are free to use, modify, and distribute this software
#          in accordance with the terms of the license.
#
# Purpose: This script provides an automated installation and management system
#          for a vulnerability assessment and penetration testing
#          toolkit. It installs and configures security tools across multiple
#          categories including exploitation, web testing, network scanning,
#          mobile security, cloud security, and Active Directory testing.
#
# DISCLAIMER: This software is provided "as-is," without warranty of any kind,
#             express or implied, including but not limited to the warranties
#             of merchantability, fitness for a particular purpose, and non-infringement.
#             In no event shall the authors or copyright holders be liable for any claim,
#             damages, or other liability, whether in an action of contract, tort, or otherwise,
#             arising from, out of, or in connection with the software or the use or other dealings
#             in the software.
#
# NOTICE: This toolkit is intended for authorized security testing only.
#         Users are responsible for ensuring compliance with all applicable laws
#         and regulations. Unauthorized use of these tools may violate local,
#         state, federal, and international laws.
#
# =============================================================================

import xml.etree.ElementTree as ET
import argparse
import sys
from pathlib import Path
from datetime import datetime
from html.parser import HTMLParser
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ZAPHTMLParser(HTMLParser):
    """Parse ZAP HTML report format"""

    def __init__(self):
        super().__init__()
        self.alerts = []
        self.current_alert = {}
        self.current_tag = None
        self.current_data = []
        self.in_alert = False
        self.in_instance = False
        self.current_instances = []

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)

        if tag == 'div' and attrs_dict.get('class') in ['alert-item', 'site']:
            self.in_alert = True
            self.current_alert = {'instances': []}
        elif tag == 'span' and 'risk-' in attrs_dict.get('class', ''):
            self.current_tag = 'risk'
        elif tag == 'h3' and self.in_alert:
            self.current_tag = 'name'
        elif tag == 'h4' and self.in_alert:
            self.current_tag = 'section'
        elif tag == 'p' and self.in_alert:
            self.current_tag = 'content'
        elif tag == 'li' and self.in_alert:
            if self.current_tag == 'urls':
                self.in_instance = True
            self.current_tag = 'list_item'
        elif tag == 'td' and self.in_alert:
            self.current_tag = 'table_cell'

    def handle_data(self, data):
        data = data.strip()
        if not data or not self.in_alert:
            return

        if self.current_tag == 'risk':
            self.current_alert['risk'] = data
        elif self.current_tag == 'name' and 'name' not in self.current_alert:
            self.current_alert['name'] = data
        elif self.current_tag == 'content':
            self.current_data.append(data)
        elif self.current_tag == 'list_item' and self.in_instance:
            self.current_instances.append(data)
        elif self.current_tag == 'table_cell':
            self.current_data.append(data)

    def handle_endtag(self, tag):
        if tag == 'div' and self.in_alert:
            if 'name' in self.current_alert and 'risk' in self.current_alert:
                content = ' '.join(self.current_data)

                if 'Description' in content:
                    parts = content.split('Description', 1)
                    if len(parts) > 1:
                        desc_part = parts[1].split('Solution', 1)[0] if 'Solution' in parts[1] else parts[1]
                        self.current_alert['description'] = desc_part.strip()

                if 'Solution' in content:
                    parts = content.split('Solution', 1)
                    if len(parts) > 1:
                        sol_part = parts[1].split('Reference', 1)[0] if 'Reference' in parts[1] else parts[1]
                        self.current_alert['solution'] = sol_part.strip()

                if 'Reference' in content:
                    parts = content.split('Reference', 1)
                    if len(parts) > 1:
                        self.current_alert['reference'] = parts[1].strip()

                if self.current_instances:
                    self.current_alert['instances'] = self.current_instances.copy()

                self.alerts.append(self.current_alert)

            self.in_alert = False
            self.current_alert = {}
            self.current_data = []
            self.current_instances = []
        elif tag in ['p', 'h3', 'h4', 'span', 'li', 'td']:
            self.current_tag = None

        if tag == 'ul':
            self.in_instance = False

def detect_format(file_path):
    """Auto-detect report type: zap_xml, zap_html, burp_xml"""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read(2000)

        if '<?xml' in content:
            if '<OWASPZAPReport' in content or 'programName="ZAP"' in content:
                return 'zap_xml'
            elif '<issues' in content and 'burpVersion' in content:
                return 'burp_xml'
            else:
                return 'unknown'
        elif '<html' in content.lower() or '<!doctype' in content.lower():
            return 'zap_html'

    return 'unknown'

def find_reports_in_directory():
    """Find all supported report files in current directory"""
    current_dir = Path('.')
    supported_files = []

    for file in current_dir.iterdir():
        if file.is_file() and file.suffix.lower() in ['.xml', '.html', '.htm']:
            try:
                report_type = detect_format(file)
                if report_type in ['zap_xml', 'zap_html', 'burp_xml']:
                    supported_files.append((file, report_type))
            except:
                continue

    return supported_files

def parse_xml_report(file_path):
    """Parse ZAP XML report"""
    tree = ET.parse(file_path)
    root = tree.getroot()

    metadata = {
        'program': root.get('programName', 'OWASP ZAP'),
        'version': root.get('version', 'Unknown'),
        'generated': root.get('generated', 'Unknown')
    }

    alerts_by_severity = defaultdict(list)
    all_alerts = []

    for site in root.findall('.//site'):
        site_name = site.get('name', 'Unknown')
        site_host = site.get('host', '')
        site_port = site.get('port', '')

        for alert in site.findall('.//alertitem'):
            alert_data = {
                'name': alert.findtext('name', 'Unknown'),
                'risk': alert.findtext('riskdesc', 'Unknown'),
                'confidence': alert.findtext('confidence', 'Unknown'),
                'description': alert.findtext('desc', ''),
                'solution': alert.findtext('solution', ''),
                'reference': alert.findtext('reference', ''),
                'cweid': alert.findtext('cweid', ''),
                'wascid': alert.findtext('wascid', ''),
                'site': site_name,
                'host': site_host,
                'port': site_port,
                'instances': []
            }

            for instance in alert.findall('.//instance'):
                instance_data = {
                    'uri': instance.findtext('uri', ''),
                    'method': instance.findtext('method', ''),
                    'param': instance.findtext('param', ''),
                    'attack': instance.findtext('attack', ''),
                    'evidence': instance.findtext('evidence', ''),
                    'request_header': instance.findtext('requestheader', ''),
                    'request_body': instance.findtext('requestbody', ''),
                    'response_header': instance.findtext('responseheader', ''),
                    'response_body': instance.findtext('responsebody', '')
                }
                alert_data['instances'].append(instance_data)

            risk_level = alert_data['risk'].split()[0] if alert_data['risk'] else 'Informational'
            alerts_by_severity[risk_level].append(alert_data)
            all_alerts.append(alert_data)

    return metadata, alerts_by_severity, all_alerts

def parse_html_report(file_path):
    """Parse ZAP HTML report"""
    with open(file_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    parser = ZAPHTMLParser()
    parser.feed(html_content)

    metadata = {
        'program': 'OWASP ZAP',
        'version': 'Unknown',
        'generated': 'Unknown'
    }

    if 'ZAP Version' in html_content:
        import re
        version_match = re.search(r'ZAP Version[:\s]+([0-9.]+)', html_content)
        if version_match:
            metadata['version'] = version_match.group(1)

    if 'Report Generated' in html_content:
        import re
        date_match = re.search(r'Report Generated[:\s]+([^<]+)', html_content)
        if date_match:
            metadata['generated'] = date_match.group(1).strip()

    alerts_by_severity = defaultdict(list)
    all_alerts = []

    for alert in parser.alerts:
        risk = alert.get('risk', 'Informational')
        if 'High' in risk:
            risk_level = 'High'
        elif 'Medium' in risk:
            risk_level = 'Medium'
        elif 'Low' in risk:
            risk_level = 'Low'
        else:
            risk_level = 'Informational'

        alert['risk'] = risk_level
        alerts_by_severity[risk_level].append(alert)
        all_alerts.append(alert)

    return metadata, alerts_by_severity, all_alerts

def parse_burp_report(file_path):
    """Parse Burp Suite XML report"""
    tree = ET.parse(file_path)
    root = tree.getroot()

    metadata = {
        'program': 'Burp Suite',
        'version': root.get('burpVersion', 'Unknown'),
        'generated': root.get('exportTime', 'Unknown')
    }

    alerts_by_severity = defaultdict(list)
    all_alerts = []

    severity_map = {
        'High': 'High',
        'Medium': 'Medium',
        'Low': 'Low',
        'Information': 'Informational',
        'Informational': 'Informational'
    }

    for issue in root.findall('.//issue'):
        alert_data = {
            'name': issue.findtext('name', 'Unknown'),
            'risk': severity_map.get(issue.findtext('severity', 'Informational'), 'Informational'),
            'confidence': issue.findtext('confidence', 'Certain'),
            'description': issue.findtext('issueBackground', '') or issue.findtext('issueDetail', ''),
            'solution': issue.findtext('remediationBackground', '') or issue.findtext('remediationDetail', ''),
            'reference': '\n'.join([ref.text for ref in issue.findall('.//reference') if ref.text]),
            'cweid': '',
            'wascid': '',
            'instances': []
        }

        host = issue.findtext('host', '')
        path = issue.findtext('path', '')

        # Extract request/response data
        request_data = issue.find('.//requestresponse/request')
        response_data = issue.find('.//requestresponse/response')

        if host and path:
            instance = {
                'uri': f"{host}{path}",
                'method': issue.findtext('method', ''),
                'param': '',
                'request_header': request_data.text if request_data is not None else '',
                'request_body': '',
                'response_header': response_data.text if response_data is not None else '',
                'response_body': '',
                'evidence': ''
            }
            alert_data['instances'].append(instance)

        risk_level = alert_data['risk']
        alerts_by_severity[risk_level].append(alert_data)
        all_alerts.append(alert_data)

    return metadata, alerts_by_severity, all_alerts

def extract_base_url(uri):
    """Extract base URL with non-standard ports"""
    from urllib.parse import urlparse

    parsed = urlparse(uri)

    # Build base URL
    base = f"{parsed.scheme}://{parsed.hostname}" if parsed.hostname else uri

    # Add port if non-standard
    if parsed.port:
        if (parsed.scheme == 'https' and parsed.port != 443) or \
           (parsed.scheme == 'http' and parsed.port != 80):
            base += f":{parsed.port}"

    return base

def add_heading(doc, text, level=1):
    """Add a heading with consistent formatting"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Calibri (Headings)'
        if level == 2:
            run.font.size = Pt(13)
        elif level == 3:
            run.font.size = Pt(11)
    return heading

def add_paragraph(doc, text, bold=False, italic=False):
    """Add a paragraph with optional formatting"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para

def add_bullet(doc, text, level=0):
    """Add a bulleted item"""
    para = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        para.paragraph_format.left_indent = Inches(0.5 * level)
    for run in para.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(11)
    return para

def generate_docx_report(metadata, alerts_by_severity, all_alerts, target_name="Target"):
    """Generate DOCX report matching nessus_parser style"""

    doc = Document()

    severity_order = ['High', 'Medium', 'Low', 'Informational']

    for severity in severity_order:
        alerts = alerts_by_severity.get(severity, [])
        if not alerts:
            continue

        # Level 2: Severity heading - Calibri (Headings) 13
        add_heading(doc, f'{severity} Severity Findings', level=2)

        # Group by alert name to avoid duplicates
        alerts_by_name = defaultdict(list)
        for alert in alerts:
            alerts_by_name[alert['name']].append(alert)

        for alert_name, alert_group in alerts_by_name.items():
            # Level 3: Finding name - Calibri (Headings) 11
            add_heading(doc, alert_name, level=3)

            alert = alert_group[0]

            # Affected Systems (first, matching nessus_parser order)
            all_instances = []
            for a in alert_group:
                all_instances.extend(a.get('instances', []))

            label = doc.add_paragraph()
            run = label.add_run('Affected System(s):')
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)

            if all_instances:
                # Extract unique base URLs
                systems = set()
                for instance in all_instances:
                    if isinstance(instance, dict) and instance.get('uri'):
                        base_url = extract_base_url(instance['uri'])
                        systems.add(base_url)

                if systems:
                    for system in sorted(systems):
                        add_bullet(doc, system)
                else:
                    add_bullet(doc, 'Unknown')
            else:
                add_bullet(doc, 'Unknown')

            # Description
            label = doc.add_paragraph()
            run = label.add_run('Description:')
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)

            if alert.get('description'):
                add_paragraph(doc, alert['description'].strip())
            else:
                p = add_paragraph(doc, 'N/A')
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Solution (called Remediation in nessus_parser)
            label = doc.add_paragraph()
            run = label.add_run('Remediation:')
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)

            if alert.get('solution'):
                add_paragraph(doc, alert['solution'].strip())
            else:
                p = add_paragraph(doc, 'N/A')
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # References
            label = doc.add_paragraph()
            run = label.add_run('References:')
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)

            if alert.get('reference'):
                has_refs = False
                for ref in alert['reference'].split('\n'):
                    ref = ref.strip()
                    if ref:
                        para = doc.add_paragraph(ref)
                        for run in para.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                        has_refs = True

                if not has_refs:
                    p = add_paragraph(doc, 'None')
                    p.runs[0].italic = True
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            else:
                p = add_paragraph(doc, 'None')
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Evidence (first instance only)
            label = doc.add_paragraph()
            run = label.add_run('Evidence:')
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(11)

            if all_instances:
                first_instance = all_instances[0]
                has_evidence = False

                # Request Header
                if first_instance.get('request_header'):
                    sub_label = doc.add_paragraph()
                    run = sub_label.add_run('Request Header:')
                    run.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)

                    para = doc.add_paragraph(first_instance['request_header'])
                    for run in para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    has_evidence = True

                # Request Body
                if first_instance.get('request_body'):
                    sub_label = doc.add_paragraph()
                    run = sub_label.add_run('Request Body:')
                    run.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)

                    para = doc.add_paragraph(first_instance['request_body'])
                    for run in para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    has_evidence = True

                # Response Header
                if first_instance.get('response_header'):
                    sub_label = doc.add_paragraph()
                    run = sub_label.add_run('Response Header:')
                    run.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)

                    para = doc.add_paragraph(first_instance['response_header'])
                    for run in para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    has_evidence = True

                # Response Body
                if first_instance.get('response_body'):
                    sub_label = doc.add_paragraph()
                    run = sub_label.add_run('Response Body:')
                    run.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)

                    response_body = first_instance['response_body']
                    if len(response_body) > 500:
                        response_body = response_body[:500] + '\n... (truncated)'

                    para = doc.add_paragraph(response_body)
                    for run in para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    has_evidence = True

                # If no evidence found, show N/A
                if not has_evidence:
                    p = add_paragraph(doc, 'N/A')
                    p.runs[0].italic = True
                    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            else:
                p = add_paragraph(doc, 'N/A')
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # Spacing between findings
            doc.add_paragraph()

    return doc

def check_and_download_hsqldb():
    """Check for HSQLDB JAR, download if missing"""
    if HSQLDB_JAR.exists():
        return True

    print(f"[!] hsqldb.jar not found in {SCRIPT_DIR}")
    response = input("Download hsqldb.jar automatically? (yes/no): ").strip().lower()

    if response != 'yes':
        print("[!] Cannot proceed without hsqldb.jar")
        print("[!] Download manually from: https://hsqldb.org/download/hsqldb_274/hsqldb.jar")
        return False

    try:
        import urllib.request
        url = "https://hsqldb.org/download/hsqldb_274/hsqldb.jar"
        print(f"[*] Downloading from {url}...")
        urllib.request.urlretrieve(url, HSQLDB_JAR)
        print(f"[+] Downloaded to {HSQLDB_JAR}")
        return True
    except Exception as e:
        print(f"[!] Download failed: {e}")
        print(f"[!] Manually download from: https://hsqldb.org/download/hsqldb_274/hsqldb.jar")
        return False

def print_test_output(metadata, alerts_by_severity, all_alerts, target_name):
    """Print test output showing document structure"""
    print("\n" + "="*60)
    print("TEST MODE - Document Structure Preview")
    print("="*60 + "\n")

    print(f"TITLE: {metadata['program']} Security Assessment Report")
    print(f"SUBTITLE: {target_name}\n")

    print("SECTION: Executive Summary")
    total_findings = len(all_alerts)
    unique_findings = len(set(alert['name'] for alert in all_alerts))
    print(f"  {total_findings} total findings, {unique_findings} unique\n")

    print("SECTION: Scan Information")
    print(f"  - Scanner: {metadata['program']} {metadata['version']}")
    print(f"  - Scan Date: {metadata['generated']}")
    print(f"  - Total: {total_findings}, Unique: {unique_findings}\n")

    print("SECTION: Findings Summary")
    severity_order = ['High', 'Medium', 'Low', 'Informational']
    for severity in severity_order:
        count = len(alerts_by_severity.get(severity, []))
        if count > 0:
            print(f"  {severity}: {count} finding(s)")
    print()

    line_count = 0
    max_lines = 30

    for severity in severity_order:
        alerts = alerts_by_severity.get(severity, [])
        if not alerts or line_count >= max_lines:
            continue

        print(f"SECTION: {severity} Severity Findings")
        line_count += 1

        alerts_by_name = defaultdict(list)
        for alert in alerts:
            alerts_by_name[alert['name']].append(alert)

        shown = 0
        for alert_name, alert_group in alerts_by_name.items():
            if shown >= 2 or line_count >= max_lines:
                remaining = len(alerts_by_name) - shown
                if remaining > 0:
                    print(f"  ... and {remaining} more {severity} findings")
                break

            print(f"  FINDING: {alert_name}")
            alert = alert_group[0]

            print(f"    - Severity: {severity}")
            if 'confidence' in alert and alert['confidence']:
                print(f"    - Confidence: {alert['confidence']}")

            instance_count = sum(len(a.get('instances', [])) for a in alert_group)
            if instance_count > 0:
                print(f"    - Affected Locations: {instance_count}")

            print()
            shown += 1
            line_count += 4

    print("="*60)
    print(f"Full report would contain all {unique_findings} unique findings")
    print("="*60)

def process_report(input_path, report_type, args):
    """Process a single report file"""
    try:
        if report_type == 'zap_xml':
            metadata, alerts_by_severity, all_alerts = parse_xml_report(input_path)
        elif report_type == 'zap_html':
            metadata, alerts_by_severity, all_alerts = parse_html_report(input_path)
        elif report_type == 'burp_xml':
            metadata, alerts_by_severity, all_alerts = parse_burp_report(input_path)

        print(f"[+] Parsed {len(all_alerts)} total findings from {input_path.name}")

    except Exception as e:
        print(f"[!] Error parsing report: {e}")
        import traceback
        traceback.print_exc()
        return

    if args.test:
        print_test_output(metadata, alerts_by_severity, all_alerts, args.target)
        return

    try:
        doc = generate_docx_report(metadata, alerts_by_severity, all_alerts, args.target)
    except Exception as e:
        print(f"[!] Error generating DOCX: {e}")
        import traceback
        traceback.print_exc()
        return

    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_name(f"{input_path.stem}_report.docx")

    try:
        doc.save(str(output_path))
        print(f"[+] Report written to: {output_path}")
    except Exception as e:
        print(f"[!] Error writing report: {e}")

def main():
    print("""
╔══════════════════════════════════════════════════════════════╗
║        ZAP Alert Selective Deletion                          ║
╚══════════════════════════════════════════════════════════════╝
    """)

    # Check for HSQLDB JAR
    if not check_and_download_hsqldb():
        return

    # Get session file
    session_name = get_session_file()
    if not session_name:
        return

    # Backup and connect
    backup_session(session_name)
    conn = connect_db(session_name)

    # Check table structure
    columns = get_table_columns(conn)
    print(f"[*] Available columns: {', '.join(columns)}")

    while True:
        # Show risk summary
        risk_summary = get_risk_summary(conn)
        risk_names = {0: 'Informational', 1: 'Low', 2: 'Medium', 3: 'High'}

        print(f"\n[*] Alerts by Risk Level:")
        for risk, count in risk_summary:
            print(f"    {risk_names.get(risk, f'Unknown({risk})')}: {count}")

        # Get and display alerts
        alerts, column_names = get_all_alerts(conn, columns)

        if not alerts:
            print("[*] No alerts in database")
            break

        display_alerts(alerts, column_names)

        print("Commands:")
        print("  - Enter numbers to delete (e.g., '1,2,3' or '1-10' or '1-5,8,10-15')")
        print("  - 'risk <level>' to delete all of a risk level (e.g., 'risk info' or 'risk low')")
        print("  - 'plugin <id>' to delete all from a plugin (e.g., 'plugin 10054')")
        print("  - 'r' to refresh list")
        print("  - 'q' to quit")

        selection = input("\nYour selection: ").strip().lower()

        if selection == 'q':
            break
        elif selection == 'r':
            continue
        elif selection.startswith('risk '):
            risk_level = selection.split(' ', 1)[1]
            delete_by_risk(conn, risk_level)
            continue
        elif selection.startswith('plugin '):
            plugin_id = selection.split(' ', 1)[1]
            delete_by_plugin(conn, plugin_id)
            continue

        # Parse selection for numeric deletions
        indices = parse_selection(selection, len(alerts))

        if not indices:
            print("[!] No valid selections")
            continue

        # Get column indices
        try:
            first_id_idx = column_names.index('FIRST_ID')
        except ValueError:
            print("[!] Unable to find alert IDs")
            continue

        plugin_idx = column_names.index('PLUGINID') if 'PLUGINID' in column_names else None
        alert_idx = column_names.index('ALERT') if 'ALERT' in column_names else None

        print(f"\n[*] Selected {len(indices)} alert type(s) for deletion:")
        for idx in indices[:10]:
            alert = alerts[idx - 1]
            info = f"  #{idx}:"
            if plugin_idx is not None:
                info += f" Plugin={alert[plugin_idx]}"
            if alert_idx is not None:
                info += f" {alert[alert_idx]}"
            print(info)

        if len(indices) > 10:
            print(f"  ... and {len(indices) - 10} more")

        confirm = input("\nDelete all instances of selected alert types? (yes/no): ").strip().lower()
        if confirm == 'yes':
            # Delete by PLUGINID and ALERT name combination
            cursor = conn.cursor()
            deleted_total = 0

            for idx in indices:
                alert = alerts[idx - 1]
                if plugin_idx is not None and alert_idx is not None:
                    plugin_id = alert[plugin_idx]
                    alert_name = alert[alert_idx]
                    cursor.execute(
                        "DELETE FROM ALERT WHERE PLUGINID = ? AND ALERT = ?",
                        [plugin_id, alert_name]
                    )
                    deleted_total += cursor.rowcount

            conn.commit()
            print(f"[+] Deleted {deleted_total} alert instances")
        else:
            print("[*] Cancelled")

    conn.close()
    print("\n[+] Done")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n[*] Interrupted")
    except Exception as e:
        print(f"[!] Error: {e}")
        import traceback
        traceback.print_exc()
