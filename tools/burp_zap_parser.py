#!/usr/bin/env python3
"""
Security Report Parser
Parses OWASP ZAP and Burp Suite HTML reports and generates a Word document
with formatted finding tables.
"""

import argparse
import re
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict
import sys


class Finding:
    """Represents a security finding"""

    def __init__(self, title, severity, description, affected_systems,
                 remediation, references, notes, tool_source):
        self.title = title
        self.severity = severity
        self.description = description
        self.affected_systems = affected_systems if isinstance(affected_systems, list) else [affected_systems]
        self.remediation = remediation
        self.references = references if isinstance(references, list) else [references]
        self.notes = notes
        self.tool_source = tool_source

    def merge_with(self, other_finding):
        """Merge another finding of the same type into this one"""
        # Add affected systems if not already present
        for system in other_finding.affected_systems:
            if system not in self.affected_systems:
                self.affected_systems.append(system)

        # DO NOT merge notes - keep only the original proof of existence

        # Merge references
        for ref in other_finding.references:
            if ref and ref not in self.references:
                self.references.append(ref)


class ZAPParser:
    """Parser for OWASP ZAP HTML reports"""

    SEVERITY_MAP = {
        'High': 'High',
        'Medium': 'Medium',
        'Low': 'Low',
        'Informational': 'Informational'
    }

    def __init__(self, html_file):
        self.html_file = html_file
        with open(html_file, 'r', encoding='utf-8') as f:
            self.soup = BeautifulSoup(f.read(), 'html.parser')

    def parse(self):
        """Parse ZAP report and return list of Finding objects"""
        findings = []

        # Find all alert sections
        alert_sections = self.soup.find_all('li', class_='alerts--site-li')

        for site_section in alert_sections:
            site_name = site_section.find('h4')
            if site_name:
                site_name = site_name.get_text(strip=True).split('(')[0].strip()

            # Find all alert types within this site - using direct children only
            alert_type_lis = []
            for child_ol in site_section.find_all('ol', recursive=False):
                for li in child_ol.find_all('li', recursive=False):
                    alert_type_lis.append(li)

            for alert_li in alert_type_lis:
                # Get alert name from h5
                h5 = alert_li.find('h5', recursive=False)
                if not h5:
                    continue

                alert_link = h5.find('a')
                if not alert_link:
                    continue

                alert_name = alert_link.get_text(strip=True)

                # Get all instances of this alert - look in nested ol > li > details
                instance_lis = []
                for ol in alert_li.find_all('ol', recursive=False):
                    instance_lis.extend(ol.find_all('li', recursive=False))

                for instance_li in instance_lis:
                    # Get the details tag
                    detail_item = instance_li.find('details', recursive=False)
                    if not detail_item:
                        continue

                    # Get URL from summary
                    summary = detail_item.find('summary')
                    url = ""
                    if summary:
                        url_span = summary.find('span', class_='request-method-n-url')
                        if url_span:
                            url = url_span.get_text(strip=True)

                    # Find the table with alert details
                    alert_table = detail_item.find('table', class_='alerts-table')
                    if not alert_table:
                        continue

                    # Extract data from table rows
                    description = ""
                    solution = ""
                    references = []
                    request = ""
                    response = ""
                    evidence = ""

                    rows = alert_table.find_all('tr')
                    for row in rows:
                        header = row.find('th')
                        if not header:
                            continue

                        header_text = header.get_text(strip=True)
                        td = row.find('td')
                        if not td:
                            continue

                        if header_text == 'Alert description':
                            description = td.get_text(strip=True)
                        elif header_text == 'Solution':
                            solution = td.get_text(strip=True)
                        elif header_text == 'Alert tags':
                            # Extract references from links
                            links = td.find_all('a')
                            for link in links:
                                href = link.get('href', '')
                                if href and ('cwe.mitre.org' in href or 'owasp.org' in href or 'zaproxy.org' in href):
                                    references.append(href)
                        elif header_text == 'Request':
                            request_details = td.find('details')
                            if request_details:
                                code = request_details.find('code')
                                if code:
                                    request = code.get_text()[:500]  # Limit length
                        elif header_text == 'Response':
                            response_details = td.find('details')
                            if response_details:
                                code = response_details.find('code')
                                if code:
                                    response = code.get_text()[:500]  # Limit length
                        elif header_text == 'Evidence':
                            evidence = td.get_text(strip=True)

                    # Determine severity from parent section
                    severity = self._extract_severity(site_section)

                    # Build notes section
                    notes = f"Tool: OWASP ZAP\nURL: {url}"
                    if evidence:
                        notes += f"\n\nEvidence: {evidence}"
                    if request:
                        notes += f"\n\nRequest Sample:\n{request}"
                    if response:
                        notes += f"\n\nResponse Sample:\n{response}"

                    finding = Finding(
                        title=alert_name,
                        severity=severity,
                        description=description,
                        affected_systems=url,
                        remediation=solution,
                        references=references,
                        notes=notes,
                        tool_source='ZAP'
                    )

                    findings.append(finding)

        return findings

    def _extract_severity(self, element):
        """Extract severity from parent sections"""
        # Look for risk level in parent sections
        parent = element
        while parent:
            if parent.name == 'li' and parent.get('id', '').startswith('alerts--risk-'):
                risk_id = parent.get('id', '')
                # Extract risk level from id (e.g., alerts--risk-2-confidence-3)
                if 'risk-3' in risk_id:
                    return 'High'
                elif 'risk-2' in risk_id:
                    return 'Medium'
                elif 'risk-1' in risk_id:
                    return 'Low'
                elif 'risk-0' in risk_id:
                    return 'Informational'
            parent = parent.parent
        return 'Medium'  # Default


class BurpParser:
    """Parser for Burp Suite HTML reports"""

    SEVERITY_MAP = {
        'High': 'High',
        'Medium': 'Medium',
        'Low': 'Low',
        'Information': 'Informational'
    }

    def __init__(self, html_file):
        self.html_file = html_file
        with open(html_file, 'r', encoding='utf-8') as f:
            self.soup = BeautifulSoup(f.read(), 'html.parser')

    def parse(self):
        """Parse Burp report and return list of Finding objects"""
        findings = []

        # Find all issue sections (marked by BODH0 class for main issues)
        issue_sections = self.soup.find_all('span', class_='BODH0')

        for issue_section in issue_sections:
            issue_link = issue_section.find('a')
            if not issue_link:
                continue

            issue_title = issue_link.get_text(strip=True)
            issue_id = issue_section.get('id', '')

            # Gather general background and remediation
            general_description = ""
            general_remediation = ""
            general_references = []

            # Find sections belonging to this issue
            next_sibling = issue_section.find_next_sibling()
            instances = []

            while next_sibling:
                if isinstance(next_sibling, str):
                    next_sibling = next_sibling.find_next_sibling()
                    continue

                if next_sibling.name == 'span' and 'BODH0' in next_sibling.get('class', []):
                    break  # reached next issue

                if next_sibling.name == 'h2':
                    header_text = next_sibling.get_text(strip=True)
                    content_elem = next_sibling.find_next_sibling()
                    if content_elem and content_elem.name == 'span':
                        content = content_elem.get_text(strip=True)
                        if header_text == 'Issue background':
                            general_description = content
                        elif header_text == 'Issue remediation':
                            general_remediation = content
                        elif header_text == 'References':
                            links = content_elem.find_all('a')
                            for link in links:
                                href = link.get('href', '')
                                if href:
                                    general_references.append(href)

                # Capture instances (sub-findings)
                if next_sibling.name == 'span' and 'BODH1' in next_sibling.get('class', []):
                    instances.append(next_sibling)

                next_sibling = next_sibling.find_next_sibling()

            # If no sub-instances, create one from summary
            if not instances:
                severity = "Medium"
                host = ""
                path = ""

                summary_table = issue_section.find_next('table', class_='summary_table')
                if summary_table:
                    rows = summary_table.find_all('tr')
                    for row in rows:
                        tds = row.find_all('td')
                        if len(tds) >= 2:
                            label = tds[0].get_text(strip=True).rstrip(':')
                            value = tds[1].get_text(strip=True)
                            if label == 'Severity':
                                severity = self.SEVERITY_MAP.get(value, value)
                            elif label == 'Host':
                                host = value
                            elif label == 'Path':
                                path = value

                affected_system = f"{host}{path}" if host else "See report for details"
                finding = Finding(
                    title=issue_title,
                    severity=severity,
                    description=general_description,
                    affected_systems=affected_system,
                    remediation=general_remediation,
                    references=general_references,
                    notes=f"Tool: Burp Suite",
                    tool_source='Burp'
                )
                findings.append(finding)
                continue

            # Process each sub-instance
            for instance in instances:
                instance_text = instance.get_text(strip=True)
                instance_url = instance_text.split('\xa0', 1)[-1] if '\xa0' in instance_text else instance_text

                summary_table = instance.find_next('table', class_='summary_table')

                severity = "Medium"
                host = ""
                path = ""

                if summary_table:
                    rows = summary_table.find_all('tr')
                    for row in rows:
                        tds = row.find_all('td')
                        if len(tds) >= 2:
                            label = tds[0].get_text(strip=True).rstrip(':')
                            value = tds[1].get_text(strip=True)
                            if label == 'Severity':
                                severity = self.SEVERITY_MAP.get(value, value)
                            elif label == 'Host':
                                host = value
                            elif label == 'Path':
                                path = value

                instance_description = general_description
                instance_detail = ""

                # Locate "Issue detail"
                next_elem = summary_table.find_next_sibling() if summary_table else instance.find_next_sibling()
                while next_elem:
                    if isinstance(next_elem, str):
                        next_elem = next_elem.find_next_sibling()
                        continue
                    if next_elem.name == 'span' and ('BODH0' in next_elem.get('class', []) or 'BODH1' in next_elem.get('class', [])):
                        break
                    if next_elem.name == 'h2' and 'Issue detail' in next_elem.get_text():
                        detail_elem = next_elem.find_next_sibling()
                        if detail_elem and detail_elem.name == 'span':
                            instance_detail = detail_elem.get_text(strip=True)[:500]
                    next_elem = next_elem.find_next_sibling()
                    if instance_detail:
                        break

                # ✅ NEW SECTION: Extract Request/Response
                request = response = ""
                rr_divs = []
                next_rr = next_elem
                while next_rr:
                    if isinstance(next_rr, str):
                        next_rr = next_rr.find_next_sibling()
                        continue
                    if next_rr.name == 'div' and 'rr_div' in next_rr.get('class', []):
                        rr_divs.append(next_rr)
                    elif next_rr.name == 'span' and ('BODH0' in next_rr.get('class', []) or 'BODH1' in next_rr.get('class', [])):
                        break
                    next_rr = next_rr.find_next_sibling()

                if len(rr_divs) >= 1:
                    request = rr_divs[0].get_text(strip=True)[:500]
                if len(rr_divs) >= 2:
                    response = rr_divs[1].get_text(strip=True)[:500]

                affected_system = f"{host}{path}" if host else instance_url

                # Build notes including request/response
                notes = f"Tool: Burp Suite\nURL: {affected_system}"
                if instance_detail:
                    notes += f"\n\nDetails:\n{instance_detail}"
                if request:
                    notes += f"\n\nRequest Sample:\n{request}"
                if response:
                    notes += f"\n\nResponse Sample:\n{response}"

                finding = Finding(
                    title=issue_title,
                    severity=severity,
                    description=instance_description if instance_description else instance_detail,
                    affected_systems=affected_system,
                    remediation=general_remediation,
                    references=general_references,
                    notes=notes,
                    tool_source='Burp'
                )

                findings.append(finding)

        return findings

class ReportGenerator:
    """Generates Word document with findings"""

    SEVERITY_ORDER = {
        'Critical': 0,
        'High': 1,
        'Medium': 2,
        'Low': 3,
        'Informational': 4
    }

    def __init__(self, findings):
        self.findings = findings
        self.doc = Document()

    def deduplicate_findings(self):
        """Deduplicate findings by title and merge affected systems"""
        finding_dict = {}

        for finding in self.findings:
            key = finding.title.lower().strip()

            if key in finding_dict:
                # Merge with existing finding
                finding_dict[key].merge_with(finding)
            else:
                finding_dict[key] = finding

        return list(finding_dict.values())

    def sort_findings(self, findings):
        """Sort findings by severity"""
        return sorted(findings,
                     key=lambda f: self.SEVERITY_ORDER.get(f.severity, 99))

    def generate(self, output_file):
        """Generate Word document with findings tables"""
        # Filter out informational findings
        filtered_findings = [f for f in self.findings if f.severity != 'Informational']

        # Store original count for reporting
        informational_count = len(self.findings) - len(filtered_findings)

        # Update findings list to filtered version
        self.findings = filtered_findings

        # Deduplicate and sort
        unique_findings = self.deduplicate_findings()
        sorted_findings = self.sort_findings(unique_findings)

        # Add title
        title = self.doc.add_heading('Security Assessment Findings', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add note
        note = self.doc.add_paragraph()
        note.add_run('Note: ').bold = True
        note.add_run('This document contains parsed findings from automated security scanning tools. ')
        note.add_run('Each finding is presented in a table format for easy integration into formal reports.')

        self.doc.add_paragraph()  # Spacing

        # Generate table for each finding
        for idx, finding in enumerate(sorted_findings, 1):
            # Add finding number
            heading = self.doc.add_heading(f'{finding.title}', level=3)

            # Create table
            table = self.doc.add_table(rows=7, cols=2)
            table.style = 'Table Grid'

            # Set column widths
            table.columns[0].width = Inches(2.0)
            table.columns[1].width = Inches(4.5)

            # Title row
            self._set_cell_content(table.cell(0, 0), 'Title of Finding', bold=True)
            self._set_cell_content(table.cell(0, 1), finding.title)

            # Severity row
            self._set_cell_content(table.cell(1, 0), 'Severity', bold=True)
            self._set_cell_content(table.cell(1, 1), finding.severity)

            # Description row
            self._set_cell_content(table.cell(2, 0), 'Description', bold=True)
            self._set_cell_content(table.cell(2, 1), finding.description)

            # System Affected row
            self._set_cell_content(table.cell(3, 0), 'System Affected', bold=True)
            systems_text = '\n'.join(finding.affected_systems) if finding.affected_systems else 'N/A'
            self._set_cell_content(table.cell(3, 1), systems_text)

            # Remediation row
            self._set_cell_content(table.cell(4, 0), 'Remediation', bold=True)
            self._set_cell_content(table.cell(4, 1), finding.remediation if finding.remediation else 'See references for remediation guidance.')

            # Reference row
            self._set_cell_content(table.cell(5, 0), 'Reference', bold=True)
            refs_text = '\n'.join(finding.references) if finding.references else 'N/A'
            self._set_cell_content(table.cell(5, 1), refs_text)

            # Notes row
            self._set_cell_content(table.cell(6, 0), 'Notes', bold=True)
            self._set_cell_content(table.cell(6, 1), finding.notes)

            # Add spacing between findings
            self.doc.add_paragraph()

        # Save document
        self.doc.save(output_file)
        print(f"\n[+] Report generated successfully: {output_file}")
        print(f"[+] Total unique findings: {len(sorted_findings)}")
        print(f"[+] Original findings before deduplication: {len(filtered_findings)}")
        if informational_count > 0:
            print(f"[+] Informational findings excluded: {informational_count}")

    def _set_cell_content(self, cell, text, bold=False):
        """Set cell content with formatting"""
        cell.text = text
        if bold:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        else:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

def find_reports_in_directory(directory):
    """
    Scan directory for ZAP and Burp HTML reports
    Returns: (list of zap files, list of burp files)
    """
    directory = Path(directory)

    if not directory.exists():
        print(f"[!] Error: Directory not found: {directory}")
        sys.exit(1)

    if not directory.is_dir():
        print(f"[!] Error: Path is not a directory: {directory}")
        sys.exit(1)

    zap_reports = []
    burp_reports = []

    # Get all HTML files in directory
    html_files = list(directory.glob("*.html")) + list(directory.glob("*.htm"))

    print(f"[*] Scanning directory: {directory}")
    print(f"[*] Found {len(html_files)} HTML file(s)")

    for html_file in html_files:
        try:
            with open(html_file, 'r', encoding='utf-8') as f:
                content = f.read(5000)  # Read first 5000 chars to identify

                # Check for ZAP signatures
                if 'ZAP by Checkmarx' in content or 'ZAP Scanning Report' in content or 'zaproxy.org' in content:
                    zap_reports.append(str(html_file))
                    print(f"  [ZAP] {html_file.name}")

                # Check for Burp signatures
                elif 'Burp Scanner Report' in content or 'Burp Suite' in content or 'portswigger' in content:
                    burp_reports.append(str(html_file))
                    print(f"  [BURP] {html_file.name}")

                else:
                    print(f"  [SKIP] {html_file.name} - Unknown format")

        except Exception as e:
            print(f"  [ERROR] {html_file.name} - Could not read: {e}")

    return zap_reports, burp_reports

def main():
    parser = argparse.ArgumentParser(
        description='Parse OWASP ZAP and Burp Suite HTML reports and generate Word document',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process all reports in a directory (auto-detects ZAP and Burp)
  python parse_security_reports.py -d ./scan_reports -o findings

  # Process specific files manually
  python parse_security_reports.py -z report1.html report2.html -o findings
  python parse_security_reports.py -b burp1.html burp2.html -o findings
  python parse_security_reports.py -z zap1.html -b burp1.html -o findings

Note: Informational findings are automatically excluded from the output.
      Duplicate findings across multiple files are merged with all affected systems listed.
        """
    )

    parser.add_argument('-d', '--directory', type=str,
                       help='Directory containing ZAP and/or Burp HTML reports (auto-detects report types)')
    parser.add_argument('-z', '--zap', type=str, nargs='+',
                       help='Path(s) to ZAP HTML report(s) - can specify multiple files')
    parser.add_argument('-b', '--burp', type=str, nargs='+',
                       help='Path(s) to Burp Suite HTML report(s) - can specify multiple files')
    parser.add_argument('-o', '--output', type=str, required=True,
                       help='Output Word document path')

    args = parser.parse_args()

    # Validate inputs - either directory OR files must be specified
    if not args.directory and not args.zap and not args.burp:
        print("[!] Error: Must specify either -d (directory) OR -z/-b (specific files)")
        parser.print_help()
        sys.exit(1)

    # Cannot use both directory and specific files
    if args.directory and (args.zap or args.burp):
        print("[!] Error: Cannot use -d (directory) with -z/-b (specific files). Choose one approach.")
        sys.exit(1)

    # Ensure output file has .docx extension
    if not args.output.lower().endswith('.docx'):
        args.output = f"{args.output}.docx"
        print(f"[*] Output filename adjusted to: {args.output}")

    # If directory mode, auto-detect reports
    if args.directory:
        zap_files, burp_files = find_reports_in_directory(args.directory)

        if not zap_files and not burp_files:
            print("[!] Error: No ZAP or Burp reports found in directory")
            sys.exit(1)

        print(f"\n[+] Identified {len(zap_files)} ZAP report(s) and {len(burp_files)} Burp report(s)")
    else:
        # Manual file mode
        zap_files = args.zap if args.zap else []
        burp_files = args.burp if args.burp else []

    all_findings = []

    # Parse ZAP report(s)
    if zap_files:
        print(f"\n[*] Processing {len(zap_files)} ZAP report(s)...")
        for zap_file in zap_files:
            if not Path(zap_file).exists():
                print(f"[!] Error: ZAP report not found: {zap_file}")
                sys.exit(1)

            print(f"  [*] Parsing: {Path(zap_file).name}")
            zap_parser = ZAPParser(zap_file)
            zap_findings = zap_parser.parse()
            all_findings.extend(zap_findings)
            print(f"      Found {len(zap_findings)} findings")

    # Parse Burp report(s)
    if burp_files:
        print(f"\n[*] Processing {len(burp_files)} Burp report(s)...")
        for burp_file in burp_files:
            if not Path(burp_file).exists():
                print(f"[!] Error: Burp report not found: {burp_file}")
                sys.exit(1)

            print(f"  [*] Parsing: {Path(burp_file).name}")
            burp_parser = BurpParser(burp_file)
            burp_findings = burp_parser.parse()
            all_findings.extend(burp_findings)
            print(f"      Found {len(burp_findings)} findings")

    if not all_findings:
        print("\n[!] Warning: No findings were extracted from the reports")
        sys.exit(0)

    # Generate report
    print(f"\n[*] Generating Word document...")
    print(f"[*] Total findings collected: {len(all_findings)}")
    generator = ReportGenerator(all_findings)
    generator.generate(args.output)

    print(f"\n[+] Done! You can now copy/paste findings from {args.output} into your formal report.")

if __name__ == '__main__':
    main()
